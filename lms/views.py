# lms/views.py
import json
import uuid
import csv
import io
import re
from datetime import datetime, timedelta
from urllib.parse import urlparse, parse_qs
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate, update_session_auth_hash
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.contrib.auth.views import LoginView, LogoutView
from django.utils import timezone
from django.utils.text import slugify
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse, Http404
from django.core.paginator import Paginator
from django.core.mail import send_mail
from django.db.models import Count, Avg, Q, Sum, F
from django.db import transaction, IntegrityError
from django.urls import reverse
from django.conf import settings
from django.template.loader import render_to_string
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.core.files.base import ContentFile
from django.core.exceptions import PermissionDenied

from .models import (
    User, Module, Course, Chapter, Quiz, Question, QuizAttempt,
    Assignment, Submission, Enrollment, EnrollmentRequest, Certificate,
    CourseReview, Notification, SystemLog, StudentProgress,
    CourseSchedule, CourseIntake, IntakeEnrollment, ClassSession, ScheduleEnrollment, SessionAttendance,
    CalendarEvent, CourseAdvertisement, CourseReminder,
    CertificationProvider, Certification, ExamRegistration, StudentCertification,
    DocumentStorage
)

# -------------------------------------------------------------------
# Helper Functions
# -------------------------------------------------------------------

def get_item(dictionary, key):
    """Template filter to get item from dictionary"""
    return dictionary.get(key, '') if dictionary else ''

def log_action(user, action, obj, request, changes=None):
    """Log user action for audit trail"""
    SystemLog.objects.create(
        user=user if user and user.is_authenticated else None,
        action=action,
        object_type=obj.__class__.__name__,
        object_id=str(obj.id),
        object_repr=str(obj)[:200],
        changes=changes or {},
        ip_address=get_client_ip(request),
        user_agent=request.META.get('HTTP_USER_AGENT', '')[:500]
    )

def get_client_ip(request):
    """Get client IP address from request"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip

def send_notification(user, title, message, notif_type='info', related_url=''):
    """Send notification to a user"""
    Notification.objects.create(
        recipient=user,
        title=title,
        message=message,
        notification_type=notif_type,
        related_url=related_url
    )

def send_email_notification(subject, message, recipient_list, html_message=None):
    """Send email notification (with fallback for development)"""
    if settings.DEBUG:
        print(f"[EMAIL PREVIEW] {subject} to {recipient_list}")
        print(f"[CONTENT PREVIEW] {message[:200]}...")
        return True
    try:
        send_mail(
            subject,
            message,
            settings.DEFAULT_FROM_EMAIL,
            recipient_list,
            html_message=html_message,
            fail_silently=False
        )
        return True
    except Exception as e:
        print(f"[EMAIL ERROR] {e}")
        return False

def render_with_ajax_support(request, template_name, context=None):
    """Render template with AJAX support - returns just content if ?ajax=true"""
    if context is None:
        context = {}
    
    if request.GET.get('ajax') == 'true':
        # For AJAX requests, render just the content without wrapping
        # Create a minimal wrapper with header and content
        return render(request, template_name, context)
    else:
        # Normal page render
        return render(request, template_name, context)


def build_embed_video_url(video_url):
    """Convert common YouTube URL formats into an embeddable URL."""
    if not video_url:
        return ""

    url = video_url.strip()
    if not url:
        return ""

    # Keep already embeddable links; ensure lightweight params.
    if "youtube.com/embed/" in url:
        return f"{url.split('?')[0]}?rel=0&modestbranding=1"

    try:
        parsed = urlparse(url)
    except Exception:
        return ""

    host = (parsed.netloc or "").lower()
    path = (parsed.path or "").strip("/")

    if "youtu.be" in host and path:
        return f"https://www.youtube.com/embed/{path}?rel=0&modestbranding=1"

    if "youtube.com" in host:
        if path == "watch":
            video_id = parse_qs(parsed.query).get("v", [""])[0]
            if video_id:
                return f"https://www.youtube.com/embed/{video_id}?rel=0&modestbranding=1"
        if path.startswith("shorts/"):
            video_id = path.split("/", 1)[1]
            if video_id:
                return f"https://www.youtube.com/embed/{video_id}?rel=0&modestbranding=1"

    # Allow direct non-YouTube embeds only when they are explicit iframe-friendly URLs.
    if url.startswith("https://") and ("vimeo.com/video/" in url or "player.vimeo.com/video/" in url):
        return url

    return ""


def strip_iframe_tags(html):
    """Remove iframe blocks from rich text content to avoid duplicate/broken embeds."""
    if not html:
        return ""
    return re.sub(r"<iframe\b[^>]*>.*?</iframe>", "", html, flags=re.IGNORECASE | re.DOTALL)

def check_role_required(roles):
    """Decorator factory for role-based access control"""
    def decorator(view_func):
        def wrapper(request, *args, **kwargs):
            if not request.user.is_authenticated:
                messages.error(request, "Please log in to access this page.")
                return redirect('lms:login')
            if request.user.role not in roles and not request.user.is_superuser:
                messages.error(request, "You don't have permission to access this page.")
                return redirect('lms:unified_dashboard')
            return view_func(request, *args, **kwargs)
        return wrapper
    return decorator

# Role-based decorators
def admin_required(view_func):
    return user_passes_test(lambda u: u.is_authenticated and (u.role == 'admin' or u.is_superuser))(view_func)

def instructor_required(view_func):
    return user_passes_test(lambda u: u.is_authenticated and u.role in ['admin', 'instructor', 'module_head'])(view_func)

def learner_required(view_func):
    return user_passes_test(lambda u: u.is_authenticated and u.role == 'learner')(view_func)

# -------------------------------------------------------------------
# Document Parsing & Question Import
# -------------------------------------------------------------------

def extract_text_from_pdf(file_obj):
    """Extract text from PDF file"""
    try:
        import PyPDF2
        text = []
        pdf_reader = PyPDF2.PdfReader(file_obj)
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)
    except ImportError:
        return "Error: PyPDF2 not installed. Please install with: pip install PyPDF2"
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file_obj):
    """Extract text from Word document"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_obj)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return '\n'.join(text)
    except ImportError:
        return "Error: python-docx not installed. Please install with: pip install python-docx"
    except Exception as e:
        return f"Error reading Word document: {str(e)}"

def extract_text_from_txt(file_obj):
    """Extract text from text file"""
    try:
        content = file_obj.read()
        if isinstance(content, bytes):
            return content.decode('utf-8', errors='ignore')
        return content
    except Exception as e:
        return f"Error reading text file: {str(e)}"

def parse_questions_from_text(text):
    """
    Parse questions from extracted text.
    Supports simple formats like:
    1. Question text here?
    a) Answer choice
    b) Answer choice
    OR
    Q1: Question text
    Correct answer: ...
    """
    questions = []
    lines = text.split('\n')
    
    current_question = None
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line starts with number followed by dot or colon (question)
        if re.match(r'^\d+[\.\)]\s+', line) or re.match(r'^Q\d+:', line):
            if current_question:
                questions.append(current_question)
            current_question = {
                'text': re.sub(r'^\d+[\.\)]\s+|^Q\d+:\s+', '', line),
                'type': 'short',
                'options': [],
                'note': ''
            }
        elif current_question and re.match(r'^[a-d][\.\)]\s+', line):
            # This is an answer option
            option_text = re.sub(r'^[a-d][\.\)]\s+', '', line)
            current_question['options'].append(option_text)
            current_question['type'] = 'mcq'
    
    if current_question:
        questions.append(current_question)
    
    return questions

def import_document_for_assignment(uploaded_file):
    """Main function to import questions from uploaded document"""
    filename = uploaded_file.name.lower()
    text = ""
    
    try:
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(uploaded_file)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(uploaded_file)
        elif filename.endswith('.txt'):
            text = extract_text_from_txt(uploaded_file)
        else:
            return {
                'success': False,
                'message': 'Unsupported file type. Please use PDF, DOCX, or TXT',
                'questions': []
            }
        
        if 'Error reading' in text:
            return {
                'success': False,
                'message': text,
                'questions': []
            }
        
        # Parse questions from text
        questions = parse_questions_from_text(text)
        
        return {
            'success': True,
            'message': f'Successfully imported {len(questions)} questions',
            'questions': questions,
            'raw_text': text[:500]  # Preview of first 500 chars
        }
    except Exception as e:
        return {
            'success': False,
            'message': f'Error processing document: {str(e)}',
            'questions': []
        }

# -------------------------------------------------------------------
# Public Views
# -------------------------------------------------------------------

def index(request):
    """Landing page"""
    from .models import CourseIntake
    
    featured_courses = Course.objects.filter(status='published', is_active=True)[:6]
    modules = Module.objects.filter(status='active')[:7]
    
    # Get upcoming course intakes
    upcoming_intakes = CourseIntake.objects.filter(
        is_visible=True,
        course_start_date__gte=timezone.now()
    ).select_related('course', 'instructor').order_by('course_start_date')[:6]
    
    # Get departments
    departments = Module.objects.filter(status='active')[:8]
    
    stats = {
        'total_students': User.objects.filter(role='learner', is_approved=True).count(),
        'total_courses': Course.objects.filter(status='published').count(),
        'total_instructors': User.objects.filter(role='instructor').count(),
        'completion_rate': 85,  # Calculate or set appropriately
    }
    context = {
        'featured_courses': featured_courses,
        'modules': modules,
        'upcoming_intakes': upcoming_intakes,
        'departments': departments,
        'stats': stats,
    }
    return render(request, 'lms/index.html', context)

def about(request):
    """About page"""
    return render(request, 'lms/about.html')

def contact(request):
    """Contact page"""
    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        message = request.POST.get('message')
        # Send email to admin
        send_email_notification(
            f"Contact Form: {name}",
            f"From: {name} ({email})\n\n{message}",
            [settings.DEFAULT_FROM_EMAIL]
        )
        messages.success(request, "Thank you for contacting us. We'll get back to you soon.")
        return redirect('lms:contact')
    return render(request, 'lms/contact.html')


def featured_courses(request):
    """Featured courses landing page with brutalism design"""
    from django.db.models import Count
    
    # Get featured/published courses with enrollment count
    courses = Course.objects.filter(
        is_published=True
    ).annotate(
        enrollment_count=Count('enrollment')
    ).order_by('-created_at')[:12]
    
    # Get statistics
    total_courses = Course.objects.filter(is_published=True).count()
    total_students = User.objects.filter(role='student').count()
    total_modules = Module.objects.count()
    
    context = {
        'featured_courses': courses,
        'total_courses': total_courses,
        'featured_users': total_students,
        'total_modules': total_modules,
    }
    
    return render(request, 'lms/landing.html', context)

# -------------------------------------------------------------------
# Authentication Views
# -------------------------------------------------------------------

class CustomLoginView(LoginView):
    template_name = 'lms/registration/login.html'
    redirect_authenticated_user = False
    
    def form_valid(self, form):
        """Record successful login attempt"""
        user = form.get_user()
        response = super().form_valid(form)
        self.request.user.record_login_attempt(success=True)
        log_action(self.request.user, 'login', self.request.user, self.request)
        
        # All users redirect to unified dashboard based on their role
        return redirect('lms:unified_dashboard')
    
    def form_invalid(self, form):
        """Record failed login attempt"""
        username = form.cleaned_data.get('username')
        if username:
            try:
                user = User.objects.get(username=username)
                user.record_login_attempt(success=False)
            except User.DoesNotExist:
                pass
        return super().form_invalid(form)

class CustomLogoutView(LogoutView):
    http_method_names = ['get', 'post', 'options']
    next_page = 'lms:index'

    def get(self, request, *args, **kwargs):
        """Support logout via GET for nav links."""
        return self.post(request, *args, **kwargs)
    
    def dispatch(self, request, *args, **kwargs):
        log_action(request.user, 'logout', request.user, request)
        return super().dispatch(request, *args, **kwargs)

def self_register(request):
    """Self-registration with approval workflow"""
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        phone_number = request.POST.get('phone_number')
        course_id = request.POST.get('course_id')
        
        # Validation
        if password != confirm_password:
            messages.error(request, "Passwords do not match.")
            return redirect('lms:register')
        
        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists.")
            return redirect('lms:register')
        
        if User.objects.filter(email=email).exists():
            messages.error(request, "Email already registered.")
            return redirect('lms:register')
        
        # Create user (pending approval)
        user = User.objects.create_user(
            username=username,
            email=email,
            password=password,
            first_name=first_name,
            last_name=last_name,
            phone_number=phone_number,
            role='learner',
            is_active=True,
            is_approved=False
        )
        
        # Generate approval token
        token = user.generate_approval_token()
        
        # Create enrollment request
        if course_id:
            course = get_object_or_404(Course, id=course_id)
            EnrollmentRequest.objects.create(
                student=user,
                course=course,
                reason=request.POST.get('reason', '')
            )
        
        # Notify admins and instructors
        approvers = User.objects.filter(
            Q(role='admin') | Q(role='approver') | Q(role='module_head')
        )
        approval_url = request.build_absolute_uri(reverse('lms:approve_registration', args=[token]))
        
        for approver in approvers:
            send_notification(
                approver,
                f"New Registration Approval Required",
                f"User {user.get_full_name()} ({user.email}) has registered and needs approval.",
                'warning',
                approval_url
            )
            
            send_email_notification(
                f"New User Registration: {user.username}",
                f"""
                A new user has registered and needs approval.
                
                Name: {user.get_full_name()}
                Email: {user.email}
                Phone: {phone_number}
                
                Approve: {approval_url}
                """,
                [approver.email]
            )
        
        messages.success(request, "Registration submitted! You will receive an email once your account is approved.")
        return redirect('lms:registration_success')
    
    courses = Course.objects.filter(status='published', is_active=True)
    return render(request, 'lms/registration/register.html', {'courses': courses})

def registration_success(request):
    """Success page after registration"""
    return render(request, 'lms/registration/registration_success.html')

def approve_registration(request, token):
    """Approve a user registration"""
    user = get_object_or_404(User, approval_token=token)
    
    if not user.is_approval_token_valid():
        messages.error(request, "Approval link has expired. Please contact support.")
        return redirect('lms:index')
    
    user.is_approved = True
    user.approval_token = None
    user.approval_token_created_at = None
    user.save()
    
    # Process pending enrollment requests
    pending_requests = EnrollmentRequest.objects.filter(student=user, status='pending')
    for req in pending_requests:
        req.approve(request.user)
    
    log_action(request.user, 'approve', user, request)
    
    # Send welcome email
    welcome_html = render_to_string('lms/registration/welcome_email.html', {'user': user})
    send_email_notification(
        f"Welcome to Synego Institute, {user.first_name}!",
        f"Your account has been approved. You can now log in at {request.build_absolute_uri(reverse('lms:login'))}",
        [user.email],
        welcome_html
    )
    
    send_notification(user, "Account Approved!", "Your account has been approved. You can now log in and start learning.", 'success')
    
    messages.success(request, f"User {user.username} has been approved.")
    return redirect('lms:unified_dashboard' if request.user.is_authenticated else 'lms:login')

def reject_registration(request, token):
    """Reject a user registration"""
    user = get_object_or_404(User, approval_token=token)
    
    user.delete()  # Remove the pending user
    
    messages.success(request, f"Registration for {user.username} has been rejected.")
    return redirect('lms:unified_dashboard' if request.user.is_authenticated else 'lms:index')

def set_password(request, token):
    """Set password for admin-created accounts"""
    try:
        user = User.objects.get(approval_token=token)
    except User.DoesNotExist:
        messages.error(request, "Invalid or expired link.")
        return redirect('lms:login')
    
    if not user.is_approval_token_valid():
        messages.error(request, "Link has expired. Please request a new one.")
        return redirect('lms:login')
    
    if request.method == 'POST':
        password = request.POST.get('password')
        confirm = request.POST.get('confirm_password')
        
        if password != confirm:
            messages.error(request, "Passwords do not match.")
            return redirect('lms:set_password', token=token)
        
        user.set_password(password)
        user.approval_token = None
        user.approval_token_created_at = None
        user.is_approved = True
        user.save()
        
        login(request, user)
        messages.success(request, "Password set successfully! Welcome to Synego Institute.")
        return redirect('lms:unified_dashboard')
    
    return render(request, 'lms/registration/set_password.html', {'token': token})

# -------------------------------------------------------------------
# Profile Views
# -------------------------------------------------------------------

@login_required
def profile_view(request):
    """View user profile"""
    enrollments = request.user.enrollments.select_related('course').all()
    certificates = request.user.certificates.select_related('course').all()
    
    context = {
        'user': request.user,
        'enrollments': enrollments,
        'certificates': certificates,
        'completion_rate': request.user.get_completion_rate(),
    }
    return render(request, 'lms/profile.html', context)

@login_required
def edit_profile(request):
    """Edit user profile"""
    if request.method == 'POST':
        request.user.first_name = request.POST.get('first_name')
        request.user.last_name = request.POST.get('last_name')
        request.user.phone_number = request.POST.get('phone_number')
        request.user.address = request.POST.get('address')
        request.user.gender = request.POST.get('gender')
        request.user.date_of_birth = request.POST.get('date_of_birth') or None
        request.user.receive_notifications = request.POST.get('receive_notifications') == 'on'
        
        if request.FILES.get('profile_picture'):
            request.user.profile_picture = request.FILES.get('profile_picture')
        
        request.user.save()
        
        messages.success(request, "Profile updated successfully!")
        return redirect('lms:profile')
    
    return render(request, 'lms/edit_profile.html', {'user': request.user})

# -------------------------------------------------------------------
# Dashboard Views (Role-based)
# -------------------------------------------------------------------

@learner_required
def student_dashboard(request):
    """Student dashboard"""
    active_enrollments = request.user.enrollments.filter(status='active').select_related('course')
    completed_courses = request.user.enrollments.filter(status='completed').select_related('course')
    
    # Calculate progress for each enrollment
    for enrollment in active_enrollments:
        enrollment.progress_percent = enrollment.get_progress_percent()
    
    # Upcoming deadlines
    upcoming_assignments = Assignment.objects.filter(
        course__in=[e.course for e in active_enrollments],
        due_date__gte=timezone.now(),
        submissions__student=request.user,
        submissions__status__in=['draft', 'returned']
    ).distinct()[:5]
    
    # Recent certificates
    recent_certificates = request.user.certificates.select_related('course').order_by('-issued_at')[:3]
    
    # Notifications
    unread_notifications = request.user.notifications.filter(is_read=False).count()
    
    # Scheduling data - Course Intakes (Scheduled courses available)
    upcoming_schedules = CourseIntake.objects.filter(
        status__in=['upcoming', 'open'],
        course_start_date__gte=timezone.now(),
        is_visible=True
    ).select_related('course').order_by('course_start_date')[:5]
    
    # Student's enrolled intake courses with sessions
    my_schedule_enrollments = IntakeEnrollment.objects.filter(
        student=request.user,
        status__in=['enrolled', 'waitlisted', 'pending']
    ).select_related('intake', 'intake__course').order_by('intake__course_start_date')[:10]
    
    # Get upcoming class sessions for this student
    my_enrolled_intakes = IntakeEnrollment.objects.filter(
        student=request.user,
        status__in=['enrolled', 'waitlisted']
    ).values_list('intake_id', flat=True)
    
    upcoming_class_sessions = ClassSession.objects.filter(
        intake_id__in=my_enrolled_intakes,
        start_time__gte=timezone.now(),
        is_cancelled=False
    ).select_related('intake', 'intake__course').order_by('start_time')[:10]
    
    context = {
        'active_enrollments': active_enrollments,
        'completed_courses': completed_courses,
        'upcoming_assignments': upcoming_assignments,
        'recent_certificates': recent_certificates,
        'unread_notifications': unread_notifications,
        'total_quiz_attempts': request.user.quiz_attempts.count(),
        'avg_quiz_score': request.user.quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0,
        'upcoming_schedules': upcoming_schedules,
        'my_schedule_enrollments': my_schedule_enrollments,
        'upcoming_class_sessions': upcoming_class_sessions,
    }
    return render(request, 'lms/student/student_dashboard.html', context)

@login_required
def instructor_dashboard(request):
    """Instructor dashboard"""
    if request.user.role not in ['instructor', 'admin', 'module_head']:
        return redirect('lms:unified_dashboard')
    
    # Courses taught by this instructor
    courses = Course.objects.filter(department=request.user.module, is_active=True)
    
    # Statistics
    total_students = Enrollment.objects.filter(course__in=courses, status='active').count()
    pending_submissions = Submission.objects.filter(
        assignment__course__in=courses,
        status='submitted'
    ).count()
    
    # Recent submissions needing grading
    recent_submissions = Submission.objects.filter(
        assignment__course__in=courses,
        status='submitted'
    ).select_related('student', 'assignment')[:10]
    
    # Course performance
    course_stats = []
    for course in courses:
        enrolled = Enrollment.objects.filter(course=course, status='active').count()
        completed = Enrollment.objects.filter(course=course, status='completed').count()
        avg_grade = Submission.objects.filter(
            assignment__course=course,
            status='graded',
            grade__isnull=False
        ).aggregate(Avg('grade'))['grade__avg'] or 0
        
        course_stats.append({
            'course': course,
            'enrolled': enrolled,
            'completed': completed,
            'completion_rate': (completed / enrolled * 100) if enrolled > 0 else 0,
            'avg_grade': avg_grade
        })
    
    # Course Intakes (Schedules)
    my_intakes = CourseIntake.objects.filter(
        instructor=request.user,
        is_visible=True
    ).prefetch_related('sessions').order_by('-course_start_date')
    
    # Upcoming class sessions
    from datetime import timedelta
    upcoming_sessions = ClassSession.objects.filter(
        intake__instructor=request.user,
        start_time__gte=timezone.now(),
        is_cancelled=False
    ).select_related('intake', 'intake__course').order_by('start_time')[:8]
    
    context = {
        'courses': courses,
        'total_students': total_students,
        'pending_submissions': pending_submissions,
        'recent_submissions': recent_submissions,
        'course_stats': course_stats,
        'my_intakes': my_intakes,
        'upcoming_sessions': upcoming_sessions,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/instructor/instructor_dashboard.html', context)

@login_required
def admin_dashboard(request):
    """Admin dashboard"""
    if request.user.role != 'admin' and not request.user.is_superuser:
        return redirect('lms:unified_dashboard')
    
    # Statistics
    total_users = User.objects.count()
    total_students = User.objects.filter(role='learner').count()
    total_instructors = User.objects.filter(role='instructor').count()
    total_courses = Course.objects.count()
    total_departments = Module.objects.count()
    pending_approvals = EnrollmentRequest.objects.filter(status='pending').count()
    pending_submissions = Submission.objects.filter(status='submitted').count()
    
    # Certification Statistics
    total_cert_providers = CertificationProvider.objects.filter(is_active=True).count()
    total_certifications = Certification.objects.filter(is_active=True).count()
    total_exam_registrations = ExamRegistration.objects.exclude(status='cancelled').count()
    
    # Schedule/Intake Statistics
    total_intakes = CourseIntake.objects.filter(is_visible=True).count()
    total_sessions = ClassSession.objects.filter(is_cancelled=False).count()
    upcoming_intakes = CourseIntake.objects.filter(
        course_start_date__gte=timezone.now(),
        is_visible=True
    ).select_related('course').order_by('course_start_date')[:5]
    today_sessions = ClassSession.objects.filter(
        start_time__date=timezone.now().date(),
        is_cancelled=False
    ).select_related('intake', 'intake__course').order_by('start_time')
    
    # Calculate average attendance across all sessions
    from django.db.models import Count, Q
    attendance_stats = SessionAttendance.objects.aggregate(
        total_records=Count('id'),
        present_count=Count('id', filter=Q(status__in=['present', 'online']))
    )
    avg_attendance = 0
    if attendance_stats['total_records'] > 0:
        avg_attendance = (attendance_stats['present_count'] / attendance_stats['total_records'] * 100)
    
    # Total students enrolled in scheduled courses
    total_scheduled_students = IntakeEnrollment.objects.count()
    
    # Recent activity
    recent_users = User.objects.order_by('-date_joined')[:5]
    recent_enrollments = Enrollment.objects.select_related('student', 'course').order_by('-enrolled_at')[:5]
    recent_cert_registrations = ExamRegistration.objects.select_related(
        'certification', 'student'
    ).order_by('-created_at')[:5]
    recent_logs = SystemLog.objects.select_related('user').order_by('-created_at')[:10]
    manageable_courses = Course.objects.filter(is_active=True).select_related('department').order_by('-updated_at')[:6]
    
    # Monthly enrollment trends (last 6 months)
    monthly_data = []
    for i in range(6):
        month_start = timezone.now().replace(day=1) - timedelta(days=30 * i)
        month_end = month_start + timedelta(days=32)
        count = Enrollment.objects.filter(enrolled_at__gte=month_start, enrolled_at__lt=month_end).count()
        monthly_data.append({
            'month': month_start.strftime('%B %Y'),
            'count': count
        })
    
    context = {
        'total_users': total_users,
        'total_students': total_students,
        'total_instructors': total_instructors,
        'total_courses': total_courses,
        'total_departments': total_departments,
        'pending_approvals': pending_approvals,
        'pending_submissions': pending_submissions,
        'total_cert_providers': total_cert_providers,
        'total_certifications': total_certifications,
        'total_exam_registrations': total_exam_registrations,
        'total_intakes': total_intakes,
        'total_sessions': total_sessions,
        'avg_attendance': round(avg_attendance, 1),
        'total_scheduled_students': total_scheduled_students,
        'upcoming_intakes': upcoming_intakes,
        'today_sessions': today_sessions,
        'recent_users': recent_users,
        'recent_enrollments': recent_enrollments,
        'recent_cert_registrations': recent_cert_registrations,
        'recent_logs': recent_logs,
        'manageable_courses': manageable_courses,
        'monthly_data': monthly_data[::-1],
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/admin/admin_dashboard.html', context)

@login_required
def module_head_dashboard(request):
    """Module head dashboard"""
    if request.user.role != 'module_head' or not request.user.module:
        return redirect('lms:unified_dashboard')
    
    module = request.user.module
    courses = module.courses.filter(is_active=True)
    instructors = module.members.filter(role='instructor')
    students = module.members.filter(role='learner', is_approved=True)
    
    context = {
        'module': module,
        'courses': courses,
        'instructors': instructors,
        'students': students,
        'total_courses': courses.count(),
        'total_instructors': instructors.count(),
        'total_students': students.count(),
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/module_head_dashboard.html', context)

# -------------------------------------------------------------------
# Unified Dashboard (Role-based Dynamic Dashboard)
# -------------------------------------------------------------------

@login_required
def unified_dashboard(request):
    """Single unified dashboard that adapts based on user role and permissions"""
    from django.db.models import Avg, Count, Q
    from django.views.decorators.http import require_http_methods
    
    user = request.user
    context = {
        'user': user,
        'user_role': user.role,
        'is_admin': user.role == 'admin' or user.is_superuser,
        'is_instructor': user.role in ['instructor', 'module_head', 'admin'],
        'is_learner': user.role == 'learner',
        'is_module_head': user.role == 'module_head',
        'unread_notifications': user.notifications.filter(is_read=False).count() if user.is_authenticated else 0,
    }
    
    # ==================== LEARNER/STUDENT CONTENT ====================
    if user.role == 'learner':
        # Active and completed courses
        active_enrollments = user.enrollments.filter(status='active').select_related('course')
        completed_courses = user.enrollments.filter(status='completed').select_related('course')
        
        # Progress calculation
        for enrollment in active_enrollments:
            enrollment.progress_percent = enrollment.get_progress_percent()
        
        # Upcoming deadlines
        upcoming_assignments = Assignment.objects.filter(
            course__in=[e.course for e in active_enrollments],
            due_date__gte=timezone.now()
        ).select_related('course').order_by('due_date')[:5]
        
        # Recent certificates
        recent_certificates = user.certificates.select_related('course').order_by('-issued_at')[:3]
        
        # Analytics
        total_quiz_attempts = user.quiz_attempts.count()
        avg_quiz_score = user.quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0
        
        # Chart data: Department breakdown
        from django.db.models import Count as DjangoCount
        dept_enrollments = Enrollment.objects.filter(
            student=user,
            status__in=['active', 'completed']
        ).values('course__department__name').annotate(
            count=DjangoCount('id')
        ).order_by('-count')[:5]
        dept_labels = [d['course__department__name'] for d in dept_enrollments]
        dept_counts = [d['count'] for d in dept_enrollments]
        
        # Quiz scores progress (last 5 attempts)
        quiz_attempts = user.quiz_attempts.order_by('attempted_at')[:5]
        quiz_labels = [f"Quiz {i+1}" for i in range(len(quiz_attempts))]
        quiz_scores = [round(q.score) for q in quiz_attempts]
        if not quiz_scores:
            quiz_scores = [0, 0, 0, 0, 0]
            quiz_labels = ['Quiz 1', 'Quiz 2', 'Quiz 3', 'Quiz 4', 'Quiz 5']
        
        # Course Intakes and Classes
        my_intakes = IntakeEnrollment.objects.filter(
            student=user,
            status__in=['enrolled', 'waitlisted']
        ).select_related('intake', 'intake__course').order_by('intake__course_start_date')[:5]
        
        upcoming_classes = ClassSession.objects.filter(
            intake__enrollments__student=user,
            start_time__gte=timezone.now(),
            is_cancelled=False
        ).select_related('intake', 'intake__course').order_by('start_time')[:5]
        
        context.update({
            'active_enrollments': active_enrollments,
            'active_enrollments_count': active_enrollments.count(),
            'completed_courses': completed_courses,
            'upcoming_assignments': upcoming_assignments,
            'recent_certificates': recent_certificates,
            'certificates_count': user.certificates.count(),
            'total_quiz_attempts': total_quiz_attempts,
            'avg_quiz_score': round(avg_quiz_score, 1),
            'pending_submissions': Submission.objects.filter(
                student=user,
                status__in=['submitted', 'grading']
            ).count(),
            'my_intakes': my_intakes,
            'upcoming_classes': upcoming_classes,
            # Chart data
            'dept_labels': dept_labels,
            'dept_counts': dept_counts,
            'quiz_labels': quiz_labels,
            'quiz_scores': quiz_scores,
        })
    
    # ==================== INSTRUCTOR CONTENT ====================
    if user.role in ['instructor', 'module_head', 'admin'] or user.is_superuser:
        # Get courses taught
        if user.role == 'admin' or user.is_superuser:
            courses = Course.objects.filter(is_active=True)
        else:
            courses = Course.objects.filter(department=user.module, is_active=True)
        
        # Statistics
        total_students = Enrollment.objects.filter(course__in=courses, status='active').count()
        pending_submissions = Submission.objects.filter(
            assignment__course__in=courses,
            status='submitted'
        ).count()
        
        # Recent submissions needing grading
        recent_submissions = Submission.objects.filter(
            assignment__course__in=courses,
            status='submitted'
        ).select_related('student', 'assignment')[:7]
        
        # Course performance stats
        course_stats = []
        for course in courses[:6]:
            enrolled = Enrollment.objects.filter(course=course, status='active').count()
            completed = Enrollment.objects.filter(course=course, status='completed').count()
            avg_grade = Submission.objects.filter(
                assignment__course=course,
                status='graded',
                grade__isnull=False
            ).aggregate(Avg('grade'))['grade__avg'] or 0
            
            course_stats.append({
                'course': course,
                'enrolled': enrolled,
                'completed': completed,
                'completion_rate': (completed / enrolled * 100) if enrolled > 0 else 0,
                'avg_grade': round(avg_grade, 1)
            })
        
        # Course Intakes
        my_intakes = CourseIntake.objects.filter(
            instructor=user,
            is_visible=True
        ).select_related('course').order_by('-course_start_date')[:8]
        
        # Upcoming class sessions
        upcoming_sessions = ClassSession.objects.filter(
            intake__instructor=user,
            start_time__gte=timezone.now(),
            is_cancelled=False
        ).select_related('intake', 'intake__course').order_by('start_time')[:8]
        
        context.update({
            'courses': courses,
            'active_enrollments_count': total_students,
            'total_students': total_students,
            'pending_submissions': pending_submissions,
            'recent_submissions': recent_submissions,
            'certificates_count': StudentCertification.objects.count(),
            'course_stats': course_stats,
            'my_intakes': my_intakes,
            'upcoming_sessions': upcoming_sessions,
        })
    
    # ==================== ADMIN CONTENT ====================
    if user.role == 'admin' or user.is_superuser:
        # System Statistics
        total_users = User.objects.count()
        total_students = User.objects.filter(role='learner').count()
        total_instructors = User.objects.filter(role='instructor').count()
        total_departments = Module.objects.count()
        pending_approvals = EnrollmentRequest.objects.filter(status='pending').count()
        pending_submissions = Submission.objects.filter(status='submitted').count()
        
        # Certification Statistics
        total_cert_providers = CertificationProvider.objects.filter(is_active=True).count()
        total_certifications = Certification.objects.filter(is_active=True).count()
        
        # Schedule Statistics
        total_intakes = CourseIntake.objects.filter(is_visible=True).count()
        total_sessions = ClassSession.objects.filter(is_cancelled=False).count()
        upcoming_intakes = CourseIntake.objects.filter(
            course_start_date__gte=timezone.now(),
            is_visible=True
        ).select_related('course').order_by('course_start_date')[:5]
        
        today_sessions = ClassSession.objects.filter(
            start_time__date=timezone.now().date(),
            is_cancelled=False
        ).select_related('intake', 'intake__course').order_by('start_time')
        
        # Attendance stats
        attendance_stats = SessionAttendance.objects.aggregate(
            total_records=Count('id'),
            present_count=Count('id', filter=Q(status__in=['present', 'online']))
        )
        avg_attendance = 0
        if attendance_stats['total_records'] > 0:
            avg_attendance = (attendance_stats['present_count'] / attendance_stats['total_records'] * 100)
        
        # Chart data: Top courses by enrollment
        from django.db.models import Count as DjangoCount
        top_courses = Course.objects.annotate(
            enrollment_count=DjangoCount('enrollments', filter=Q(enrollments__status='active'))
        ).order_by('-enrollment_count')[:5]
        course_labels = [c.title for c in top_courses]
        course_counts = [c.enrollment_count for c in top_courses]
        if not course_counts:
            course_labels = ['Course 1', 'Course 2', 'Course 3', 'Course 4', 'Course 5']
            course_counts = [0, 0, 0, 0, 0]
        
        # Chart data: User growth (last 6 months)
        from datetime import timedelta
        user_growth_labels = []
        user_growth_counts = []
        for i in range(5, -1, -1):
            month_date = timezone.now() - timedelta(days=30*i)
            month_start = month_date.replace(day=1)
            if i > 0:
                month_end = (month_date.replace(day=1) + timedelta(days=32)).replace(day=1) - timedelta(days=1)
            else:
                month_end = timezone.now()
            
            users_in_month = User.objects.filter(
                date_joined__gte=month_start,
                date_joined__lte=month_end
            ).count()
            user_growth_labels.append(month_date.strftime('%b'))
            user_growth_counts.append(users_in_month)
        
        # Recent logs and users
        recent_logs = SystemLog.objects.select_related('user').order_by('-created_at')[:10] if hasattr(SystemLog, 'objects') else []
        recent_users = User.objects.order_by('-date_joined')[:5]
        recent_enrollments = Enrollment.objects.select_related('student', 'course').order_by('-enrolled_at')[:5]
        
        context.update({
            'active_enrollments_count': Enrollment.objects.filter(status='active').count(),
            'total_users': total_users,
            'total_students': total_students,
            'total_instructors': total_instructors,
            'total_departments': total_departments,
            'pending_approvals': pending_approvals,
            'pending_submissions': Submission.objects.filter(status='submitted').count(),
            'total_quiz_attempts': QuizAttempt.objects.count(),
            'total_cert_providers': total_cert_providers,
            'total_certifications': total_certifications,
            'certificates_count': StudentCertification.objects.count(),
            'total_intakes': total_intakes,
            'total_sessions': total_sessions,
            'avg_attendance': round(avg_attendance, 1),
            'upcoming_intakes': upcoming_intakes,
            'today_sessions': today_sessions,
            'recent_logs': recent_logs,
            'recent_users': recent_users,
            'recent_enrollments': recent_enrollments,
            # Chart data
            'course_labels': course_labels,
            'course_counts': course_counts,
            'user_growth_labels': user_growth_labels,
            'user_growth_counts': user_growth_counts,
        })
    
    # ==================== MODULE HEAD CONTENT ====================
    if user.role == 'module_head' and user.module:
        module = user.module
        module_courses = module.courses.filter(is_active=True)
        module_instructors = module.members.filter(role='instructor')
        module_students = module.members.filter(role='learner', is_approved=True)
        
        context.update({
            'module': module,
            'module_courses': module_courses,
            'module_instructors': module_instructors,
            'module_students': module_students,
            'total_module_courses': module_courses.count(),
            'total_module_instructors': module_instructors.count(),
            'total_module_students': module_students.count(),
        })
    
    return render(request, 'lms/dashboard.html', context)

# -------------------------------------------------------------------
# Module Views
# -------------------------------------------------------------------

def module_list(request):
    """List all modules"""
    modules = Module.objects.filter(status='active')
    return render(request, 'lms/modules/module_list.html', {'modules': modules})


def module_detail(request, slug):
    """Module detail page"""
    module = get_object_or_404(Module, slug=slug, status='active')
    courses = list(module.courses.filter(status='published', is_active=True).order_by('title'))

    level_map = {
        'beginner': 'Certificate / Short Course',
        'intermediate': 'Certificate / Diploma',
        'advanced': 'Diploma / Advanced Certificate',
    }
    for course in courses:
        course.programme_level = level_map.get(course.difficulty, course.get_difficulty_display())

    sample_curriculum = []
    if not courses and 'civil' in module.name.lower():
        sample_curriculum = [
            {
                'programme': 'Engineering Surveying',
                'level': 'Certificate / Diploma',
                'duration': '6 - 12 months',
            },
            {
                'programme': 'Road Construction and Maintenance',
                'level': 'Certificate / Diploma',
                'duration': '6 - 12 months',
            },
            {
                'programme': 'Concrete Technology and Testing',
                'level': 'Certificate / Short Course',
                'duration': '3 months',
            },
            {
                'programme': 'Structural Design Fundamentals',
                'level': 'Certificate',
                'duration': '6 months',
            },
            {
                'programme': 'Materials Testing Laboratory Practice',
                'level': 'Short Course',
                'duration': '4 - 8 weeks',
            },
            {
                'programme': 'Environmental Compliance in Construction',
                'level': 'Short Course',
                'duration': '2 weeks',
            },
        ]

    return render(request, 'lms/modules/module_detail.html', {
        'module': module,
        'courses': courses,
        'sample_curriculum': sample_curriculum,
        'total_programmes': len(courses) + len(sample_curriculum),
    })

# -------------------------------------------------------------------
# Course Views
# -------------------------------------------------------------------

def course_list(request):
    """List all courses with filtering and pagination"""
    courses = Course.objects.filter(status='published', is_active=True).select_related('department')
    
    # Filtering
    module_param = request.GET.get('module')
    if module_param:
        # Handle both module ID and slug
        if module_param.isdigit():
            courses = courses.filter(department_id=module_param)
        else:
            courses = courses.filter(department__slug=module_param)
    
    difficulty = request.GET.get('difficulty')
    if difficulty:
        courses = courses.filter(difficulty=difficulty)
    
    search = request.GET.get('q')
    if search:
        courses = courses.filter(Q(title__icontains=search) | Q(description__icontains=search))
    
    # Pagination
    paginator = Paginator(courses, 12)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    departments = Module.objects.filter(status='active')
    
    # Featured Scheduled Courses
    from lms.models.scheduling import CourseSchedule
    featured_schedules = CourseSchedule.objects.filter(
        status__in=['open', 'upcoming'],
        is_featured=True,
        is_visible=True,
        course_start_date__gte=timezone.now()
    ).select_related('course').order_by('course_start_date')[:6]
    
    context = {
        'page_obj': page_obj,
        'modules': departments,
        'selected_module': module_param,
        'selected_difficulty': difficulty,
        'search_query': search,
        'featured_schedules': featured_schedules,
        'now': timezone.now(),
    }
    return render(request, 'lms/courses/course_list.html', context)

@learner_required
def my_enrolled_courses(request):
    """List courses the learner is enrolled in"""
    enrollments = request.user.enrollments.filter(status__in=['active', 'completed']).select_related('course')
    courses = [e.course for e in enrollments]
    
    # Pagination
    paginator = Paginator(courses, 12)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'page_title': 'My Enrolled Courses',
        'is_my_courses': True,
    }
    return render(request, 'lms/courses/course_list.html', context)

def course_redirect_to_slug(request, course_id):
    """Redirect old numeric course IDs to slug-based URLs"""
    course = get_object_or_404(Course, id=course_id, is_active=True)
    return redirect('lms:course_detail', slug=course.slug, permanent=True)

def course_detail(request, slug):
    """Course detail page"""
    course = get_object_or_404(Course, slug=slug, is_active=True)
    
    # Check if user is enrolled
    is_enrolled = False
    enrollment = None
    progress = {}
    
    if request.user.is_authenticated and request.user.role == 'learner':
        enrollment = Enrollment.objects.filter(student=request.user, course=course).first()
        is_enrolled = enrollment is not None and enrollment.status == 'active'
        
        if is_enrolled:
            progress = enrollment.progress
            # Calculate overall progress
            total_chapters = course.chapters.count()
            completed_chapters = len(progress.get('completed_chapters', []))
            enrollment.progress_percent = (completed_chapters / total_chapters * 100) if total_chapters > 0 else 0
    
    # Get chapters
    chapters = course.chapters.all().order_by('order')
    
    # Check if there's a pending enrollment request
    has_pending_request = False
    if request.user.is_authenticated:
        has_pending_request = EnrollmentRequest.objects.filter(
            student=request.user, course=course, status='pending'
        ).exists()
    
    # Get scheduled offerings for this course (CourseIntakes)
    from lms.models.scheduling import CourseIntake
    
    now = timezone.now()
    upcoming_intakes = CourseIntake.objects.filter(
        course=course,
        is_visible=True,
        course_start_date__gte=now
    ).select_related('instructor').order_by('course_start_date')
    
    past_intakes = CourseIntake.objects.filter(
        course=course,
        is_visible=True,
        course_end_date__lt=now
    ).select_related('instructor').order_by('-course_end_date')
    
    context = {
        'course': course,
        'chapters': chapters,
        'is_enrolled': is_enrolled,
        'enrollment': enrollment,
        'has_pending_request': has_pending_request,
        'total_chapters': chapters.count(),
        'upcoming_intakes': upcoming_intakes,
        'past_intakes': past_intakes,
        'now': now,
    }
    return render(request, 'lms/courses/course_detail.html', context)

@login_required
def enroll_course(request, slug):
    """Enroll in a course (direct enrollment or request)"""
    course = get_object_or_404(Course, slug=slug, is_active=True)
    
    # Check if already enrolled
    if Enrollment.objects.filter(student=request.user, course=course).exists():
        messages.warning(request, "You are already enrolled in this course.")
        return redirect('lms:course_detail', slug=course.slug)
    
    # Check if there's a pending request
    if EnrollmentRequest.objects.filter(student=request.user, course=course, status='pending').exists():
        messages.warning(request, "You already have a pending enrollment request.")
        return redirect('lms:course_detail', slug=course.slug)
    
    # If course requires approval or user is self-registered, create request
    enrollment_request = EnrollmentRequest.objects.create(
        student=request.user,
        course=course,
        reason=request.POST.get('reason', 'Direct enrollment request')
    )
    
    # Notify instructors
    instructors = User.objects.filter(
        Q(role='instructor', department=course.department) |
        Q(role='admin') |
        Q(role='approver')
    )
    
    approval_url = request.build_absolute_uri(reverse('lms:admin_enrollment_requests'))
    for instructor in instructors:
        send_notification(
            instructor,
            f"New Enrollment Request: {course.title}",
            f"{request.user.get_full_name()} has requested to enroll in {course.title}.",
            'info',
            approval_url
        )
    
    messages.success(request, "Your enrollment request has been submitted. You will be notified when approved.")
    return redirect('lms:course_detail', slug=course.slug)

@login_required
def unenroll_course(request, slug):
    """Unenroll from a course"""
    course = get_object_or_404(Course, slug=slug)
    enrollment = get_object_or_404(Enrollment, student=request.user, course=course)
    
    enrollment.status = 'dropped'
    enrollment.save()
    
    log_action(request.user, 'update', enrollment, request, {'status': 'dropped'})
    messages.success(request, f"You have been unenrolled from {course.title}.")
    return redirect('lms:course_list')

# -------------------------------------------------------------------
# Chapter Views
# -------------------------------------------------------------------

@login_required
def chapter_detail(request, course_slug, chapter_id):
    """View a chapter's content"""
    course = get_object_or_404(Course, slug=course_slug, is_active=True)
    chapter = get_object_or_404(Chapter, id=chapter_id, course=course)
    
    # Check access
    is_admin_access = request.user.is_superuser or request.user.role == 'admin'
    enrollment = Enrollment.objects.filter(student=request.user, course=course).first()
    
    if not is_admin_access and not chapter.is_free_preview and (not enrollment or enrollment.status != 'active'):
        messages.error(request, "Please enroll in this course to access this chapter.")
        return redirect('lms:course_detail', slug=course.slug)
    
    # Check if previous chapters are completed (if required)
    if not is_admin_access and chapter.requires_previous_quiz_pass and not chapter.is_free_preview:
        previous_chapters = Chapter.objects.filter(course=course, order__lt=chapter.order)
        progress = enrollment.progress if enrollment else {}
        completed_chapters = progress.get('completed_chapters', [])
        
        for prev in previous_chapters:
            if str(prev.id) not in completed_chapters:
                messages.warning(request, f"Please complete Chapter {prev.order} first.")
                return redirect('lms:course_detail', slug=course.slug)
    
    # Get quiz for this chapter (if exists)
    quiz = None
    if hasattr(chapter, 'quiz'):
        quiz = chapter.quiz
    
    # Get previous and next chapters
    prev_chapter = Chapter.objects.filter(course=course, order=chapter.order - 1).first()
    next_chapter = Chapter.objects.filter(course=course, order=chapter.order + 1).first()

    chapter_video_embed_url = build_embed_video_url(chapter.video_url)
    chapter_content_html = strip_iframe_tags(chapter.content)
    
    context = {
        'course': course,
        'chapter': chapter,
        'quiz': quiz,
        'prev_chapter': prev_chapter,
        'next_chapter': next_chapter,
        'enrollment': enrollment,
        'is_admin_access': is_admin_access,
        'chapter_video_embed_url': chapter_video_embed_url,
        'chapter_content_html': chapter_content_html,
    }
    return render(request, 'lms/courses/chapter_detail.html', context)

@login_required
def mark_chapter_complete(request, course_slug, chapter_id):
    """Mark a chapter as completed"""
    course = get_object_or_404(Course, slug=course_slug)
    chapter = get_object_or_404(Chapter, id=chapter_id, course=course)
    enrollment = get_object_or_404(Enrollment, student=request.user, course=course)
    
    # Update progress
    progress = enrollment.progress or {}
    completed_chapters = progress.get('completed_chapters', [])
    
    if str(chapter.id) not in completed_chapters:
        completed_chapters.append(str(chapter.id))
        progress['completed_chapters'] = completed_chapters
        enrollment.progress = progress
        enrollment.save()
        
        # Update StudentProgress cache
        StudentProgress.objects.update_or_create(
            student=request.user,
            course=course,
            defaults={'overall_percent': enrollment.get_progress_percent()}
        )
        
        messages.success(request, f"Chapter '{chapter.title}' completed!")
    
    return redirect('lms:chapter_detail', course_slug=course_slug, chapter_id=chapter_id)

# -------------------------------------------------------------------
# Quiz Views (Non-graded)
# -------------------------------------------------------------------

@login_required
def take_quiz(request, quiz_id):
    """Take a quiz (non-graded, performance appraisal only)"""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    chapter = quiz.chapter
    course = chapter.course
    
    # Check enrollment
    enrollment = get_object_or_404(Enrollment, student=request.user, course=course)
    
    # Check if already attempted and no retakes allowed
    existing_attempt = QuizAttempt.objects.filter(student=request.user, quiz=quiz).first()
    if existing_attempt and existing_attempt.passed and quiz.attempts_allowed == 1:
        messages.warning(request, "You have already passed this quiz and cannot retake it.")
        return redirect('lms:chapter_detail', course_slug=course.slug, chapter_id=chapter.id)
    
    if existing_attempt and QuizAttempt.objects.filter(student=request.user, quiz=quiz).count() >= quiz.attempts_allowed:
        messages.warning(request, "You have used all allowed attempts for this quiz.")
        return redirect('lms:chapter_detail', course_slug=course.slug, chapter_id=chapter.id)
    
    questions = quiz.questions.all()
    
    if request.method == 'POST':
        # Process answers
        score = 0
        total_points = 0
        answers = {}
        
        for question in questions:
            total_points += question.points
            user_answer = request.POST.get(f'question_{question.id}', '')
            answers[str(question.id)] = user_answer
            
            if question.check_answer(user_answer):
                score += question.points
        
        score_percent = (score / total_points * 100) if total_points > 0 else 0
        passed = score_percent >= quiz.pass_score
        
        # Record attempt
        attempt = QuizAttempt.objects.create(
            student=request.user,
            quiz=quiz,
            score=score_percent,
            passed=passed,
            answers=answers,
            completed_at=timezone.now()
        )
        
        log_action(request.user, 'create', attempt, request)
        
        # Update chapter progress if passed
        if passed:
            progress = enrollment.progress or {}
            quiz_passes = progress.get('quiz_passes', [])
            if str(quiz.id) not in quiz_passes:
                quiz_passes.append(str(quiz.id))
                progress['quiz_passes'] = quiz_passes
                enrollment.progress = progress
                enrollment.save()
        
        return redirect('lms:quiz_result', quiz_id=quiz.id)
    
    context = {
        'quiz': quiz,
        'chapter': chapter,
        'course': course,
        'questions': questions,
        'time_limit': quiz.time_limit_minutes,
    }
    return render(request, 'lms/courses/quiz_take.html', context)

@login_required
def quiz_result(request, quiz_id):
    """Show quiz results"""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    attempt = get_object_or_404(QuizAttempt, student=request.user, quiz=quiz)
    
    context = {
        'quiz': quiz,
        'attempt': attempt,
        'chapter': quiz.chapter,
    }
    return render(request, 'lms/courses/quiz_result.html', context)

@login_required
def retake_quiz(request, quiz_id):
    """Allow retaking a quiz"""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    
    # Check if retake is allowed
    attempt_count = QuizAttempt.objects.filter(student=request.user, quiz=quiz).count()
    if attempt_count >= quiz.attempts_allowed:
        messages.error(request, "You have reached the maximum number of attempts.")
        return redirect('lms:chapter_detail', course_slug=quiz.chapter.course.slug, chapter_id=quiz.chapter.id)
    
    return redirect('lms:take_quiz', quiz_id=quiz_id)

# -------------------------------------------------------------------
# Assignment Views (Student)
# -------------------------------------------------------------------

@login_required
def student_assignment_list(request):
    """List all assignments for student's enrolled courses"""
    # Get all courses the student is enrolled in
    enrollments = Enrollment.objects.filter(
        student=request.user, 
        status='active'
    ).select_related('course').values_list('course', flat=True)
    
    # Get all assignments for those courses, ordered by due date
    assignments = Assignment.objects.filter(
        course__id__in=enrollments
    ).select_related('course').order_by('due_date')
    
    # Get submission status for each assignment
    assignment_list = []
    for assignment in assignments:
        submission = Submission.objects.filter(
            assignment=assignment,
            student=request.user
        ).first()
        
        assignment_list.append({
            'assignment': assignment,
            'submission': submission,
            'is_overdue': assignment.is_overdue,
            'submission_status': submission.status if submission else 'not_started',
            'grade': submission.grade if submission else None,
        })
    
    # Pagination
    paginator = Paginator(assignment_list, 10)
    page = request.GET.get('page', 1)
    assignments_page = paginator.get_page(page)
    
    context = {
        'assignments': assignments_page,
        'total_assignments': len(assignment_list),
        'page_title': 'My Assignments',
        'total_not_started': sum(1 for a in assignment_list if a['submission_status'] == 'not_started'),
        'total_draft': sum(1 for a in assignment_list if a['submission_status'] == 'draft'),
        'total_submitted': sum(1 for a in assignment_list if a['submission_status'] == 'submitted'),
        'total_graded': sum(1 for a in assignment_list if a['submission_status'] == 'graded'),
    }
    
    return render(request, 'lms/assignments/student_assignment_list.html', context)

@login_required
def assignment_detail(request, assignment_id):
    """View assignment details - accessible to students (enrolled), instructors (teaching course), and admins"""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    course = assignment.course
    
    # Check permissions - allow instructors, admins, or enrolled students
    is_admin = request.user.is_superuser or request.user.role == 'admin'
    
    # Check if user is an instructor for any intake of this course
    is_instructor = is_admin or CourseIntake.objects.filter(
        course=course, 
        instructor=request.user
    ).exists()
    
    is_enrolled = Enrollment.objects.filter(student=request.user, course=course, status='active').exists()
    
    if not (is_instructor or is_admin or is_enrolled):
        messages.error(request, 'You do not have access to this assignment')
        return redirect('lms:course_list')
    
    # Get or create submission (only for students)
    submission = None
    if is_enrolled:
        submission, created = Submission.objects.get_or_create(
            assignment=assignment,
            student=request.user,
            defaults={'status': 'draft'}
        )
    
    # Check if assignment is overdue
    is_overdue = assignment.is_overdue
    
    # Get related documents
    related_documents = assignment.stored_documents.all()
    
    context = {
        'assignment': assignment,
        'course': course,
        'submission': submission,
        'is_overdue': is_overdue,
        'related_documents': related_documents,
        'is_instructor': is_instructor,
        'is_admin': is_admin,
    }
    return render(request, 'lms/assignments/assignments_detail.html', context)

@login_required
def submit_assignment(request, assignment_id):
    """Submit an assignment"""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    course = assignment.course
    enrollment = get_object_or_404(Enrollment, student=request.user, course=course)
    
    submission, created = Submission.objects.get_or_create(
        assignment=assignment,
        student=request.user,
        defaults={'status': 'draft'}
    )
    
    if request.method == 'POST':
        # Handle file upload
        if request.FILES.get('submitted_file'):
            submission.submitted_file = request.FILES['submitted_file']
        
        # Handle Google Doc ID
        google_doc_id = request.POST.get('google_doc_id')
        if google_doc_id:
            submission.google_doc_id = google_doc_id
        
        submission.status = 'submitted'
        submission.submitted_at = timezone.now()
        submission.ip_address = get_client_ip(request)
        submission.user_agent = request.META.get('HTTP_USER_AGENT', '')[:500]
        submission.save()
        
        # Submit to Turnitin if enabled
        if assignment.enable_plagiarism_check:
            submission.submit_to_turnitin()
        
        log_action(request.user, 'submit', submission, request)
        
        # Notify instructors for this course
        course_instructors = CourseIntake.objects.filter(
            course=course,
            status='active'
        ).values_list('instructor', flat=True)
        
        instructors = User.objects.filter(
            Q(id__in=course_instructors) |
            Q(role='admin')
        )
        grade_url = request.build_absolute_uri(reverse('lms:grade_submission', args=[submission.id]))
        
        for instructor in instructors:
            send_notification(
                instructor,
                f"New Submission: {assignment.title}",
                f"{request.user.get_full_name()} has submitted {assignment.title}",
                'info',
                grade_url
            )
        
        messages.success(request, "Assignment submitted successfully!")
        return redirect('lms:course_detail', slug=course.slug)
    
    return redirect('lms:assignment_detail', assignment_id=assignment_id)

@login_required
def view_submission(request, submission_id):
    """View a submission"""
    submission = get_object_or_404(Submission, id=submission_id, student=request.user)
    
    context = {
        'submission': submission,
        'assignment': submission.assignment,
        'course': submission.assignment.course,
    }
    return render(request, 'lms/assignments/submission_view.html', context)

@login_required
def edit_submission(request, submission_id):
    """Edit a draft submission"""
    submission = get_object_or_404(Submission, id=submission_id, student=request.user, status='draft')
    
    if request.method == 'POST':
        if request.FILES.get('submitted_file'):
            submission.submitted_file = request.FILES['submitted_file']
        if request.POST.get('google_doc_id'):
            submission.google_doc_id = request.POST.get('google_doc_id')
        submission.save()
        messages.success(request, "Submission updated.")
        return redirect('lms:assignment_detail', assignment_id=submission.assignment.id)
    
    return redirect('lms:assignment_detail', assignment_id=submission.assignment.id)

# -------------------------------------------------------------------
# Grading Views (Instructor)
# -------------------------------------------------------------------

@instructor_required
def instructor_assignments(request):
    """List all assignments for instructor"""
    # Get courses - admins/superusers see all, instructors see only their department courses
    if request.user.is_superuser or request.user.role == 'admin':
        courses = Course.objects.filter(is_active=True)
    else:
        courses = Course.objects.filter(department=request.user.module, is_active=True)
    
    # Get assignments for those courses
    assignments = Assignment.objects.filter(course__in=courses).select_related('course').order_by('-due_date')
    
    # Filtering by course (optional)
    course_slug = request.GET.get('course', '')
    if course_slug:
        assignments = assignments.filter(course__slug=course_slug)
    
    # Filtering by status
    status_filter = request.GET.get('status', '')
    if status_filter == 'open':
        assignments = assignments.filter(due_date__gte=timezone.now())
    elif status_filter == 'closed':
        assignments = assignments.filter(due_date__lt=timezone.now())
    
    # Pagination
    paginator = Paginator(assignments, 10)  # 10 assignments per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Add submission counts and analytics data
    assignment_data = []
    for assignment in page_obj:
        submission_count = assignment.submissions.count()
        graded_count = assignment.submissions.filter(status='graded').count()
        pending_count = assignment.submissions.filter(status='submitted').count()
        
        assignment_data.append({
            'assignment': assignment,
            'submission_count': submission_count,
            'graded_count': graded_count,
            'pending_count': pending_count,
            'is_closed': assignment.due_date < timezone.now(),
            'is_due_soon': timezone.now() < assignment.due_date < timezone.now() + timedelta(days=1),
        })
    
    context = {
        'assignments': assignments,
        'assignment_data': assignment_data,
        'page_obj': page_obj,
        'courses': courses,
        'selected_course': course_slug,
        'selected_status': status_filter,
        'now': timezone.now(),
    }
    return render(request, 'lms/instructor/assignments_list.html', context)

@instructor_required
def submission_list(request, assignment_id):
    """List all submissions for an assignment"""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    submissions = assignment.submissions.select_related('student').all()
    
    context = {
        'assignment': assignment,
        'submissions': submissions,
    }
    return render(request, 'lms/instructor/submission_list.html', context)

@instructor_required
def grade_submission(request, submission_id):
    """Grade a submission"""
    submission = get_object_or_404(Submission, id=submission_id)
    
    if request.method == 'POST':
        grade = float(request.POST.get('grade', 0))
        feedback = request.POST.get('feedback', '')
        
        # Calculate rubric scores if provided
        rubric_scores = {}
        if request.POST.get('rubric_scores'):
            rubric_scores = json.loads(request.POST.get('rubric_scores'))
        
        submission.grade = grade
        submission.feedback = feedback
        submission.rubric_scores = rubric_scores
        submission.status = 'graded'
        submission.graded_by = request.user
        submission.graded_at = timezone.now()
        submission.save()
        
        # Update grade in Google Classroom if integrated
        if submission.assignment.google_classroom_id:
            # This would call Google Classroom API
            pass
        
        log_action(request.user, 'grade', submission, request, {'grade': grade})
        
        # Notify student
        send_notification(
            submission.student,
            f"Grade Posted: {submission.assignment.title}",
            f"You received {grade}% on {submission.assignment.title}. Feedback: {feedback[:100]}",
            'grade',
            reverse('lms:view_submission', args=[submission.id])
        )
        
        messages.success(request, f"Submission graded: {grade}%")
        return redirect('lms:submission_list', assignment_id=submission.assignment.id)
    
    context = {
        'submission': submission,
        'assignment': submission.assignment,
        'student': submission.student,
    }
    return render(request, 'lms/instructor/grade_submission.html', context)

@instructor_required
def return_submission(request, submission_id):
    """Return a submission for revision"""
    submission = get_object_or_404(Submission, id=submission_id)
    
    submission.status = 'returned'
    submission.revision_count += 1
    submission.save()
    
    send_notification(
        submission.student,
        f"Assignment Returned: {submission.assignment.title}",
        f"Your submission has been returned for revision. Please review the feedback and resubmit.",
        'warning',
        reverse('lms:assignment_detail', args=[submission.assignment.id])
    )
    
    messages.success(request, "Submission returned to student for revision.")
    return redirect('lms:submission_list', assignment_id=submission.assignment.id)

@instructor_required
def assignment_analytics(request, assignment_id):
    """View analytics for an assignment"""
    if request.user.role not in ['instructor', 'admin', 'module_head']:
        return redirect('lms:unified_dashboard')
    
    assignment = get_object_or_404(Assignment, id=assignment_id)
    submissions = assignment.submissions.filter(status='graded', grade__isnull=False)
    
    grades = [s.grade for s in submissions]
    avg_grade = sum(grades) / len(grades) if grades else 0
    min_grade = min(grades) if grades else 0
    max_grade = max(grades) if grades else 0
    
    grade_distribution = {
        '90-100': len([g for g in grades if g >= 90]),
        '80-89': len([g for g in grades if 80 <= g < 90]),
        '70-79': len([g for g in grades if 70 <= g < 80]),
        '60-69': len([g for g in grades if 60 <= g < 70]),
        'Below 60': len([g for g in grades if g < 60]),
    }
    
    context = {
        'assignment': assignment,
        'submissions_count': submissions.count(),
        'avg_grade': avg_grade,
        'min_grade': min_grade,
        'max_grade': max_grade,
        'grade_distribution': grade_distribution,
    }
    return render(request, 'lms/instructor/assignment_analytics.html', context)

# -------------------------------------------------------------------
# Assignment Creation Views (Instructor)
# -------------------------------------------------------------------

@login_required
def api_course_chapters(request, course_slug):
    """API endpoint to fetch chapters for a course"""
    try:
        course = get_object_or_404(Course, slug=course_slug)
        
        # Check if user is an instructor in this course's department
        if request.user.role not in ['instructor', 'admin', 'module_head']:
            return JsonResponse({'error': 'Unauthorized'}, status=403)
        
        # Allow admin/superuser to access any course
        # For instructors/module_head, check if they belong to the course's module
        if not (request.user.is_superuser or request.user.role == 'admin'):
            if request.user.role in ['instructor', 'module_head'] and course.department != request.user.module:
                return JsonResponse({'error': 'Unauthorized'}, status=403)
        
        chapters = course.chapters.all().values('id', 'title', 'order')
        return JsonResponse({
            'success': True,
            'chapters': list(chapters)
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@instructor_required
@instructor_required
def create_assignment(request, course_slug=None):
    """Create a new assignment"""
    # Get all courses for the instructor
    # If user is admin/superuser, show all courses
    # Otherwise, show only courses from the user's module
    if request.user.is_superuser or request.user.role == 'admin':
        instructor_courses = Course.objects.filter(is_active=True).order_by('title')
    else:
        instructor_courses = Course.objects.filter(department=request.user.module, is_active=True).order_by('title')
    
    # Get selected course (from POST or URL parameter)
    selected_course_slug = request.POST.get('course_slug') or course_slug
    
    if not selected_course_slug and instructor_courses.exists():
        selected_course_slug = instructor_courses.first().slug
    
    # Get the course object
    if selected_course_slug:
        try:
            course = Course.objects.get(slug=selected_course_slug, is_active=True)
        except Course.DoesNotExist:
            course = None
            messages.error(request, "Course not found or is inactive")
    else:
        course = None
    
    # Handle document import - directly save without preview
    if request.method == 'POST' and 'import_document' in request.FILES:
        uploaded_file = request.FILES['import_document']
        print(f"[DEBUG] Processing import for file: {uploaded_file.name}")
        import_result = import_document_for_assignment(uploaded_file)
        
        if import_result['success']:
            # Directly save to DocumentStorage
            imported_questions = import_result['questions']
            print(f"[DEBUG] Successfully imported {len(imported_questions)} questions from {uploaded_file.name}")
            
            # Store in session for later attachment to assignment
            request.session['imported_document_file'] = uploaded_file.name
            request.session['imported_questions_data'] = imported_questions
            messages.success(request, f"✓ Document imported successfully! ({len(imported_questions)} questions found)")
            print(f"[DEBUG] Stored {len(imported_questions)} questions in session")
        else:
            messages.error(request, import_result['message'])
            print(f"[DEBUG] Import failed: {import_result['message']}")
    
    # Handle assignment creation
    if request.method == 'POST' and 'title' in request.POST and course:
        print(f"[DEBUG] Creating assignment: {request.POST.get('title')}")
        print(f"[DEBUG] Course: {course.title}")
        print(f"[DEBUG] Due Date: {request.POST.get('due_date')}")
        # Create the assignment
        imported_questions = request.session.get('imported_questions_data', [])
        print(f"[DEBUG] Found {len(imported_questions)} imported questions in session")
        
        # Get chapter if provided
        chapter = None
        chapter_id = request.POST.get('chapter_id')
        if chapter_id:
            try:
                chapter = Chapter.objects.get(id=chapter_id, course=course)
            except Chapter.DoesNotExist:
                chapter = None
        
        # Parse dates properly
        due_date = request.POST.get('due_date')
        soft_deadline = request.POST.get('soft_deadline')
        soft_deadline = soft_deadline if soft_deadline else None
        
        try:
            assignment = Assignment.objects.create(
                course=course,
                chapter=chapter,
                title=request.POST.get('title'),
                description=request.POST.get('description'),
                assignment_type=request.POST.get('assignment_type', 'individual'),
                due_date=due_date,
                soft_deadline=soft_deadline,
                late_penalty_percent=int(request.POST.get('late_penalty_percent', 10)),
                total_points=int(request.POST.get('total_points', 100)),
                rubric=json.loads(request.POST.get('rubric', '{}')),
                enable_plagiarism_check=request.POST.get('enable_plagiarism_check') == 'on',
                max_file_size_mb=int(request.POST.get('max_file_size_mb', 50)),
                allowed_file_types=request.POST.get('allowed_file_types', '.pdf,.doc,.docx,.txt'),
            )
            print(f"[DEBUG] ✓ Assignment created successfully with ID: {assignment.id}")
        except Exception as e:
            print(f"[DEBUG] ✗ Assignment creation failed: {str(e)}")
            messages.error(request, f"Error creating assignment: {str(e)}")
            return render(request, 'lms/instructor/assignment_form.html', {
                'course': course,
                'instructor_courses': Course.objects.filter(is_active=True).order_by('title'),
                'chapters': course.chapters.all() if course else [],
            })
        
        # Auto-save any imported document to DocumentStorage
        if imported_questions:
            try:
                # Create a text file from imported questions
                questions_text = "IMPORTED ASSIGNMENT QUESTIONS\n"
                questions_text += f"=" * 50 + "\n\n"
                for i, q in enumerate(imported_questions, 1):
                    questions_text += f"Question {i}: {q.get('text', '')}\n"
                    if q.get('options'):
                        questions_text += f"Options: {', '.join(q.get('options', []))}\n"
                    questions_text += "\n"
                
                # Create a file-like object from the text
                from django.core.files.base import ContentFile
                file_content = ContentFile(questions_text.encode('utf-8'), name=f"{assignment.title}_questions.txt")
                
                document = DocumentStorage.objects.create(
                    name=f"{assignment.title} - Questions",
                    description=f"Imported questions for assignment: {assignment.title}. Contains {len(imported_questions)} questions.",
                    document_type='assignment',
                    file=file_content,
                    uploaded_by=request.user,
                    related_assignment=assignment,
                    course=course,
                    is_public=True,  # Make visible to all enrolled students and course instructors
                    can_download=True,
                    can_view_online=True,
                )
                print(f"[DEBUG] ✓ Document saved to storage with ID: {document.id}")
                messages.success(request, f"Document with {len(imported_questions)} questions saved to storage!")
            except Exception as e:
                # Log the error but don't fail the assignment creation
                print(f"[DEBUG] Warning: Could not save document to storage: {str(e)}")
        
        # Clear imported questions from session after creating assignment
        if 'imported_questions_data' in request.session:
            del request.session['imported_questions_data']
        if 'imported_document_file' in request.session:
            del request.session['imported_document_file']
        
        log_action(request.user, 'create', assignment, request)
        messages.success(request, f"Assignment '{assignment.title}' created successfully!")
        return redirect('lms:instructor_assignments')
    
    chapters = course.chapters.all() if course else []
    
    # Define assignment type choices
    assignment_types = [
        ('individual', 'Individual Assignment'),
        ('group', 'Group Assignment'),
    ]
    
    # Define file type presets
    file_type_presets = {
        'documents': '.pdf,.doc,.docx,.txt,.rtf',
        'images': '.jpg,.jpeg,.png,.gif,.webp',
        'media': '.mp4,.mkv,.avi,.mov,.webm,.mp3,.wav',
        'code': '.py,.java,.js,.cpp,.c,.html,.css,.txt',
        'spreadsheet': '.xlsx,.xls,.csv,.ods',
        'all': '.pdf,.doc,.docx,.txt,.jpg,.jpeg,.png,.gif,.mp4,.mkv,.avi,.mov,.zip,.rar',
    }
    
    # Load previously imported questions from session if they exist
    imported_questions_preview = []
    if 'imported_questions_data' in request.session:
        imported_questions_preview = request.session.get('imported_questions_data', [])
        print(f"[DEBUG] Loaded {len(imported_questions_preview)} questions from session")
    
    print(f"[DEBUG] Rendering form")
    
    context = {
        'course': course,
        'instructor_courses': instructor_courses,
        'chapters': chapters,
        'assignment_types': assignment_types,
        'file_type_presets': file_type_presets,
    }
    return render(request, 'lms/instructor/assignment_form.html', context)

# -------------------------------------------------------------------
# Document Storage & Management Views
# -------------------------------------------------------------------

@login_required
def document_storage_list(request):
    """List all stored documents - visible based on user role and course enrollment"""
    if request.user.role == 'learner':
        # Learners can see: public documents + documents from their enrolled courses
        documents = DocumentStorage.objects.filter(
            Q(is_public=True) | Q(course__in=request.user.enrollments.values('course'))
        ).distinct()
    elif request.user.is_superuser or request.user.role == 'admin':
        # Admins and superusers see all documents
        documents = DocumentStorage.objects.all()
    else:
        # Instructors see: 
        # - Documents from courses they teach
        # - Documents they uploaded
        # - Public documents
        instructor_courses = Course.objects.filter(
            intakes__instructor=request.user
        ).distinct()
        documents = DocumentStorage.objects.filter(
            Q(course__in=instructor_courses) |  # Documents from their courses
            Q(uploaded_by=request.user) |        # Documents they uploaded
            Q(is_public=True)                    # Public documents
        ).distinct()
    
    # Check if filtering for "my documents"
    my_documents = request.GET.get('my', '')
    if my_documents:
        documents = documents.filter(uploaded_by=request.user)
    
    # Filtering
    doc_type = request.GET.get('type', '')
    search_query = request.GET.get('search', '')
    
    if doc_type:
        documents = documents.filter(document_type=doc_type)
    
    if search_query:
        documents = documents.filter(
            Q(name__icontains=search_query) |
            Q(description__icontains=search_query) |
            Q(tags__icontains=search_query)
        )
    
    # Pagination
    paginator = Paginator(documents.order_by('-uploaded_at'), 15)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)
    
    # Get document types for filter
    doc_types = DocumentStorage.DOCUMENT_TYPE_CHOICES
    
    context = {
        'page_obj': page_obj,
        'documents': page_obj.object_list,
        'doc_types': doc_types,
        'selected_type': doc_type,
        'search_query': search_query,
        'my_documents': my_documents,
    }
    return render(request, 'lms/document_storage/document_list.html', context)

@login_required
def document_upload(request):
    """Upload a new document"""
    if request.method == 'POST':
        name = request.POST.get('name', '')
        description = request.POST.get('description', '')
        document_type = request.POST.get('document_type', 'other')
        tags = request.POST.get('tags', '')
        is_public = request.POST.get('is_public') == 'on'
        file = request.FILES.get('file')
        course_id = request.POST.get('course')
        
        if not file or not name:
            messages.error(request, 'Please provide document name and file')
            return redirect('lms:document_storage_list')
        
        try:
            document = DocumentStorage.objects.create(
                name=name,
                description=description,
                document_type=document_type,
                tags=tags,
                is_public=is_public,
                file=file,
                uploaded_by=request.user,
                course_id=course_id if course_id else None,
            )
            messages.success(request, f'Document "{name}" uploaded successfully!')
            log_action(request.user, 'create', document, request)
        except Exception as e:
            messages.error(request, f'Error uploading document: {str(e)}')
        
        return redirect('lms:document_storage_list')
    
    # GET request - show upload form
    courses = Course.objects.filter(is_active=True)
    if not (request.user.is_superuser or request.user.role == 'admin'):
        courses = courses.filter(department=request.user.module)
    
    context = {
        'doc_types': DocumentStorage.DOCUMENT_TYPE_CHOICES,
        'courses': courses,
    }
    return render(request, 'lms/document_storage/document_upload.html', context)

@login_required
def document_detail(request, slug):
    """View document details"""
    document = get_object_or_404(DocumentStorage, slug=slug)
    
    # Check access
    if not document.can_user_access(request.user):
        messages.error(request, 'You do not have permission to access this document')
        return redirect('lms:document_storage_list')
    
    # Record view
    document.record_view()
    
    # Process tags
    tags_list = [tag.strip() for tag in document.tags.split(',') if tag.strip()] if document.tags else []
    
    context = {
        'document': document,
        'tags_list': tags_list,
        'can_download': document.can_download and document.can_user_access(request.user),
        'can_delete': document.uploaded_by == request.user or request.user.is_superuser or request.user.role == 'admin',
    }
    return render(request, 'lms/document_storage/document_detail.html', context)

@login_required
def document_download(request, slug):
    """Download document"""
    document = get_object_or_404(DocumentStorage, slug=slug)
    
    # Check access
    if not document.can_download or not document.can_user_access(request.user):
        messages.error(request, 'You do not have permission to download this document')
        return redirect('lms:document_storage_list')
    
    # Record download
    document.record_download()
    
    # Serve file
    file_path = document.file.path
    if os.path.exists(file_path):
        with open(file_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/octet-stream')
            response['Content-Disposition'] = f'attachment; filename="{document.file.name}"'
            return response
    else:
        messages.error(request, 'File not found')
        return redirect('lms:document_storage_list')

@login_required
def document_delete(request, slug):
    """Delete a document"""
    document = get_object_or_404(DocumentStorage, slug=slug)
    
    # Check permission
    if document.uploaded_by != request.user and not request.user.is_superuser and request.user.role != 'admin':
        messages.error(request, 'You do not have permission to delete this document')
        return redirect('lms:document_storage_list')
    
    if request.method == 'POST':
        doc_name = document.name
        try:
            # Delete file
            if document.file:
                document.file.delete()
            # Delete record
            document.delete()
            messages.success(request, f'Document "{doc_name}" deleted successfully')
            log_action(request.user, 'delete', document, request)
        except Exception as e:
            messages.error(request, f'Error deleting document: {str(e)}')
        
        return redirect('lms:document_storage_list')
    
    context = {'document': document}
    return render(request, 'lms/document_storage/document_delete.html', context)

@instructor_required
def document_edit(request, slug):
    """Edit document metadata"""
    document = get_object_or_404(DocumentStorage, slug=slug)
    
    # Check permission
    if document.uploaded_by != request.user and not request.user.is_superuser and request.user.role != 'admin':
        messages.error(request, 'You do not have permission to edit this document')
        return redirect('lms:document_storage_list')
    
    if request.method == 'POST':
        document.name = request.POST.get('name', document.name)
        document.description = request.POST.get('description', '')
        document.document_type = request.POST.get('document_type', 'other')
        document.tags = request.POST.get('tags', '')
        document.is_public = request.POST.get('is_public') == 'on'
        document.can_download = request.POST.get('can_download') == 'on'
        document.can_view_online = request.POST.get('can_view_online') == 'on'
        
        try:
            document.save()
            messages.success(request, 'Document updated successfully')
            log_action(request.user, 'update', document, request)
        except Exception as e:
            messages.error(request, f'Error updating document: {str(e)}')
        
        return redirect('lms:document_detail', slug=document.slug)
    
    context = {
        'document': document,
        'doc_types': DocumentStorage.DOCUMENT_TYPE_CHOICES,
    }
    return render(request, 'lms/document_storage/document_edit.html', context)

@instructor_required
def document_share(request, slug):
    """Share document with users (AJAX endpoint)"""
    document = get_object_or_404(DocumentStorage, slug=slug)
    
    # Check permission
    if document.uploaded_by != request.user and not request.user.is_superuser and request.user.role != 'admin':
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    if request.method == 'POST' and request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        share_with = request.POST.getlist('share_with[]')
        
        # TODO: Implement sharing logic with a Sharing model
        # For now, just mark as public if necessary
        
        return JsonResponse({
            'success': True,
            'message': f'Document shared with {len(share_with)} user(s)'
        })
    
    return JsonResponse({'error': 'Invalid request'}, status=400)

# Helper import for document views
import os

@instructor_required
def document_import_assignment_endpoint(request, assignment_id):
    """Import questions/document from assignment creation"""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    
    # Check if user can access this assignment
    if assignment.course.department != request.user.module and not (request.user.is_superuser or request.user.role == 'admin'):
        return redirect('lms:instructor_assignments')
    
    if request.method == 'POST' and 'import_document' in request.FILES:
        uploaded_file = request.FILES['import_document']
        
        try:
            # Create document in central storage
            document = DocumentStorage.objects.create(
                name=f"{assignment.title} - {uploaded_file.name}",
                description=f"Document imported for assignment: {assignment.title}",
                document_type='assignment',
                file=uploaded_file,
                uploaded_by=request.user,
                related_assignment=assignment,
                course=assignment.course,
                is_public=False,
            )
            messages.success(request, f'Document imported and stored successfully as "{document.name}"')
            log_action(request.user, 'create', document, request)
            return redirect('lms:document_detail', slug=document.slug)
        except Exception as e:
            messages.error(request, f'Error importing document: {str(e)}')
    
    return redirect('lms:instructor_assignments')

def store_document_automatically(file_obj, name, description, doc_type, user, course=None, assignment=None):
    """Helper function to automatically store documents from any feature"""
    try:
        document = DocumentStorage.objects.create(
            name=name,
            description=description,
            document_type=doc_type,
            file=file_obj,
            uploaded_by=user,
            related_assignment=assignment,
            course=course,
            is_public=False,
            can_download=True,
            can_view_online=True,
        )
        return document
    except Exception as e:
        print(f"Error storing document: {str(e)}")
        return None

@instructor_required
def edit_assignment(request, assignment_id):
    """Edit an existing assignment"""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    
    # Check permissions - only instructor/admin who owns the course can edit
    if request.user.role == 'instructor' and assignment.course.department != request.user.module:
        messages.error(request, "You don't have permission to edit this assignment.")
        return redirect('lms:instructor_assignments')
    
    if request.method == 'POST':
        assignment.title = request.POST.get('title')
        assignment.description = request.POST.get('description')
        assignment.assignment_type = request.POST.get('assignment_type')
        assignment.due_date = request.POST.get('due_date')
        assignment.soft_deadline = request.POST.get('soft_deadline') or None
        assignment.late_penalty_percent = int(request.POST.get('late_penalty_percent', 10))
        assignment.total_points = int(request.POST.get('total_points', 100))
        assignment.max_file_size_mb = int(request.POST.get('max_file_size_mb', 50))
        assignment.allowed_file_types = request.POST.get('allowed_file_types', '')
        assignment.rubric = json.loads(request.POST.get('rubric', '{}'))
        assignment.enable_plagiarism_check = request.POST.get('enable_plagiarism_check') == 'on'
        
        # Handle chapter change if provided
        chapter_id = request.POST.get('chapter_id')
        if chapter_id:
            try:
                chapter = Chapter.objects.get(id=chapter_id)
                assignment.chapter = chapter
            except Chapter.DoesNotExist:
                pass
        
        assignment.save()
        
        log_action(request.user, 'update', assignment, request)
        messages.success(request, "Assignment updated successfully!")
        return redirect('lms:instructor_assignments')
    
    # Get instructor's courses
    if request.user.is_superuser or request.user.role == 'admin':
        instructor_courses = Course.objects.filter(is_active=True).order_by('title')
    else:
        instructor_courses = Course.objects.filter(
            department=request.user.module,
            is_active=True
        ).order_by('title')
    
    # Get chapters for the assignment's course
    chapters = assignment.course.chapters.all().order_by('order')
    
    # Assignment types
    assignment_types = Assignment._meta.get_field('assignment_type').choices
    
    context = {
        'assignment': assignment,
        'course': assignment.course,
        'instructor_courses': instructor_courses,
        'chapters': chapters,
        'assignment_types': assignment_types,
        'page_title': f'Edit Assignment - {assignment.title}',
    }
    return render(request, 'lms/instructor/assignment_form.html', context)

@instructor_required
def delete_assignment(request, assignment_id):
    """Delete an assignment"""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    assignment.delete()
    
    log_action(request.user, 'delete', assignment, request)
    messages.success(request, "Assignment deleted.")
    return redirect('lms:instructor_assignments')


# -------------------------------------------------------------------
# Course Material Management Views (Admin & Instructor)
# -------------------------------------------------------------------

def _redirect_course_management(course_slug):
    """Redirect helper that falls back to course detail if manage route is unavailable."""
    try:
        return redirect('lms:manage_course', course_slug=course_slug)
    except Exception:
        return redirect('lms:course_detail', slug=course_slug)


@instructor_required
def manage_course(request, course_slug):
    """Course management hub for admin/instructor/department head."""
    course = get_object_or_404(Course, slug=course_slug)

    if request.user.role in ['instructor', 'module_head'] and course.department != request.user.module:
        messages.error(request, "You can only manage courses in your department.")
        return redirect('lms:course_detail', slug=course.slug)

    chapters = course.chapters.all().order_by('order')
    assignments = course.assignments.all().order_by('due_date')
    quizzes = Quiz.objects.filter(chapter__course=course).select_related('chapter').order_by('chapter__order')

    context = {
        'course': course,
        'chapters': chapters,
        'assignments': assignments,
        'quizzes': quizzes,
    }
    return render(request, 'lms/admin/manage_course.html', context)


@instructor_required
def add_course_material(request, course_slug):
    """Add material (chapter, assignment, quiz) to a course."""
    course = get_object_or_404(Course, slug=course_slug)

    if request.user.role in ['instructor', 'module_head'] and course.department != request.user.module:
        messages.error(request, "You can only add materials to courses in your module.")
        return redirect('lms:course_detail', slug=course.slug)

    chapters = course.chapters.all().order_by('order')
    context = {
        'course': course,
        'chapters': chapters,
    }
    return render(request, 'lms/admin/add_course_material.html', context)


@login_required
def select_course_for_chapter(request):
    """Select a course to add a chapter to."""
    # Check if user is admin or instructor
    is_admin = request.user.is_superuser or request.user.groups.filter(name__in=['admin']).exists()
    is_instructor = request.user.groups.filter(name__in=['instructor']).exists()
    
    if not (is_admin or is_instructor):
        raise PermissionDenied("Only admins and instructors can add chapters.")
    
    # Get courses available to the user
    if is_admin:
        # Admin: show all courses
        courses = Course.objects.all().order_by('title')
    else:
        # Instructor: show only their courses
        courses = Course.objects.filter(instructor=request.user).order_by('title')
    
    context = {
        'courses': courses,
        'page_title': 'Select Course to Add Chapter'
    }
    return render(request, 'lms/admin/select_course_for_chapter.html', context)


@admin_required
def admin_add_chapter(request, course_slug):
    """Admin view to add a chapter to a course."""
    course = get_object_or_404(Course, slug=course_slug)

    if request.method == 'POST':
        # Handle file upload - documents are primary
        document_file = None
        if 'document_file' in request.FILES:
            uploaded_file = request.FILES['document_file']
            # Validate file extension
            allowed_extensions = ('.pdf', '.pptx', '.ppt', '.xlsx', '.xls')
            if not any(uploaded_file.name.lower().endswith(ext) for ext in allowed_extensions):
                messages.error(request, "Only PDF, PowerPoint (.pptx, .ppt), and Excel (.xlsx, .xls) files are allowed.")
                return redirect('lms:admin_add_chapter', course_slug=course.slug)
            document_file = uploaded_file
        
        chapter = Chapter.objects.create(
            course=course,
            title=request.POST.get('title', '').strip(),
            order=int(request.POST.get('order', 1)),
            chapter_type=request.POST.get('chapter_type', 'lesson'),
            document_file=document_file,
            video_url=request.POST.get('video_url', '').strip(),  # Optional
            content=request.POST.get('content', ''),
            estimated_minutes=int(request.POST.get('estimated_minutes', 30)),
            template_doc_id=request.POST.get('template_doc_id', '').strip(),
            is_free_preview=request.POST.get('is_free_preview') == 'on',
            requires_previous_quiz_pass=request.POST.get('requires_previous_quiz_pass') == 'on',
        )
        log_action(request.user, 'create', chapter, request)
        messages.success(request, f"Chapter '{chapter.title}' added to {course.title}.")

        if chapter.chapter_type == 'quiz':
            return redirect('lms:admin_add_quiz_with_chapter', course_slug=course.slug, chapter_id=chapter.id)

        return _redirect_course_management(course.slug)

    context = {
        'course': course,
        'chapter_types': Chapter.CHAPTER_TYPE_CHOICES,
    }
    return render(request, 'lms/admin/add_chapter.html', context)


@admin_required
def admin_add_assignment(request, course_slug):
    """Admin view to add an assignment to a course."""
    course = get_object_or_404(Course, slug=course_slug)

    if request.method == 'POST':
        rubric_raw = request.POST.get('rubric', '{}')
        try:
            rubric_value = json.loads(rubric_raw)
        except json.JSONDecodeError:
            messages.error(request, "Rubric JSON is invalid.")
            return redirect('lms:admin_add_assignment', course_slug=course.slug)

        assignment = Assignment.objects.create(
            course=course,
            chapter_id=request.POST.get('chapter_id') or None,
            title=request.POST.get('title', '').strip(),
            description=request.POST.get('description', ''),
            assignment_type=request.POST.get('assignment_type', 'individual'),
            due_date=request.POST.get('due_date'),
            soft_deadline=request.POST.get('soft_deadline') or None,
            late_penalty_percent=int(request.POST.get('late_penalty_percent', 10)),
            total_points=int(request.POST.get('total_points', 100)),
            rubric=rubric_value,
            enable_plagiarism_check=request.POST.get('enable_plagiarism_check') == 'on',
            max_file_size_mb=int(request.POST.get('max_file_size_mb', 50)),
            allowed_file_types=request.POST.get('allowed_file_types', '.pdf,.doc,.docx,.txt'),
        )
        log_action(request.user, 'create', assignment, request)
        messages.success(request, f"Assignment '{assignment.title}' added to {course.title}.")
        return _redirect_course_management(course.slug)

    chapters = course.chapters.all().order_by('order')
    context = {
        'course': course,
        'chapters': chapters,
        'assignment_types': Assignment.ASSIGNMENT_TYPE_CHOICES,
    }
    return render(request, 'lms/admin/add_assignment.html', context)


@admin_required
def admin_add_quiz(request, course_slug, chapter_id=None):
    """Admin view to add a quiz to a chapter."""
    course = get_object_or_404(Course, slug=course_slug)
    selected_chapter = None
    if chapter_id is not None:
        selected_chapter = get_object_or_404(Chapter, id=chapter_id, course=course)

    if request.method == 'POST':
        chapter = get_object_or_404(Chapter, id=request.POST.get('chapter_id'), course=course)

        if hasattr(chapter, 'quiz'):
            messages.warning(request, f"Chapter '{chapter.title}' already has a quiz. You can edit it instead.")
            return redirect('lms:admin_edit_quiz', quiz_id=chapter.quiz.id)

        quiz = Quiz.objects.create(
            chapter=chapter,
            title=request.POST.get('title', '').strip() or f"Quiz for {chapter.title}",
            description=request.POST.get('description', ''),
            pass_score=int(request.POST.get('pass_score', 70)),
            time_limit_minutes=int(request.POST.get('time_limit_minutes')) if request.POST.get('time_limit_minutes') else None,
            attempts_allowed=int(request.POST.get('attempts_allowed', 1)),
            shuffle_questions=request.POST.get('shuffle_questions') == 'on',
            show_answers_after_submit=request.POST.get('show_answers_after_submit') == 'on',
        )
        log_action(request.user, 'create', quiz, request)
        messages.success(request, f"Quiz added to chapter '{chapter.title}'.")
        return redirect('lms:admin_add_quiz_questions', quiz_id=quiz.id)

    chapters = course.chapters.all().order_by('order')
    context = {
        'course': course,
        'chapters': chapters,
        'selected_chapter': selected_chapter,
    }
    return render(request, 'lms/admin/add_quiz.html', context)


@admin_required
def admin_add_quiz_questions(request, quiz_id):
    """Add questions to a quiz."""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    chapter = quiz.chapter
    course = chapter.course

    if request.method == 'POST':
        question_text = request.POST.get('question_text', '').strip()
        question_type = request.POST.get('question_type', 'mcq')
        points = int(request.POST.get('points', 1))
        next_order = quiz.questions.count() + 1

        if not question_text:
            messages.error(request, "Question text is required.")
            return redirect('lms:admin_add_quiz_questions', quiz_id=quiz.id)

        if question_type == 'mcq':
            options = [o.strip() for o in request.POST.getlist('options[]') if o.strip()]
            correct_answer = request.POST.get('correct_answer', '').strip()
            Question.objects.create(
                quiz=quiz,
                text=question_text,
                question_type='mcq',
                options=options,
                correct_answer=correct_answer,
                points=points,
                order=next_order,
            )
        elif question_type == 'tf':
            correct_answer = request.POST.get('correct_answer', 'True')
            Question.objects.create(
                quiz=quiz,
                text=question_text,
                question_type='tf',
                options=['True', 'False'],
                correct_answer=correct_answer,
                points=points,
                order=next_order,
            )
        elif question_type == 'short':
            correct_answer = request.POST.get('correct_answer', '').strip()
            Question.objects.create(
                quiz=quiz,
                text=question_text,
                question_type='short',
                options=[],
                correct_answer=correct_answer,
                points=points,
                order=next_order,
            )

        messages.success(request, "Question added successfully.")
        if request.POST.get('add_another') == 'on':
            return redirect('lms:admin_add_quiz_questions', quiz_id=quiz.id)
        return _redirect_course_management(course.slug)

    questions = quiz.questions.all().order_by('order', 'id')
    context = {
        'quiz': quiz,
        'chapter': chapter,
        'course': course,
        'questions': questions,
        'question_types': [
            ('mcq', 'Multiple Choice'),
            ('tf', 'True/False'),
            ('short', 'Short Answer'),
        ],
    }
    return render(request, 'lms/admin/add_quiz_questions.html', context)


@admin_required
def admin_edit_chapter(request, chapter_id):
    """Edit an existing chapter."""
    chapter = get_object_or_404(Chapter, id=chapter_id)
    course = chapter.course

    if request.method == 'POST':
        chapter.title = request.POST.get('title', '').strip()
        chapter.order = int(request.POST.get('order', 1))
        chapter.chapter_type = request.POST.get('chapter_type', chapter.chapter_type)
        chapter.content = request.POST.get('content', '')
        chapter.estimated_minutes = int(request.POST.get('estimated_minutes', 30))
        chapter.template_doc_id = request.POST.get('template_doc_id', '').strip()
        chapter.is_free_preview = request.POST.get('is_free_preview') == 'on'
        chapter.requires_previous_quiz_pass = request.POST.get('requires_previous_quiz_pass') == 'on'
        
        # Handle document file upload
        if 'document_file' in request.FILES:
            uploaded_file = request.FILES['document_file']
            # Validate file extension
            allowed_extensions = ('.pdf', '.pptx', '.ppt', '.xlsx', '.xls')
            if not any(uploaded_file.name.lower().endswith(ext) for ext in allowed_extensions):
                messages.error(request, "Only PDF, PowerPoint (.pptx, .ppt), and Excel (.xlsx, .xls) files are allowed.")
                return redirect('lms:admin_edit_chapter', chapter_id=chapter.id)
            # Delete old file if exists
            if chapter.document_file:
                chapter.document_file.delete()
            chapter.document_file = uploaded_file
        
        # Video URL is optional
        chapter.video_url = request.POST.get('video_url', '').strip()
        
        chapter.save()

        log_action(request.user, 'update', chapter, request)
        messages.success(request, f"Chapter '{chapter.title}' updated successfully.")
        return _redirect_course_management(course.slug)

    context = {
        'chapter': chapter,
        'course': course,
        'chapter_types': Chapter.CHAPTER_TYPE_CHOICES,
    }
    return render(request, 'lms/admin/edit_chapter.html', context)


@admin_required
def admin_delete_chapter(request, chapter_id):
    """Delete a chapter."""
    chapter = get_object_or_404(Chapter, id=chapter_id)
    course = chapter.course
    chapter_title = chapter.title

    if request.method == 'POST':
        log_action(request.user, 'delete', chapter, request)
        chapter.delete()
        messages.success(request, f"Chapter '{chapter_title}' deleted successfully.")
        return _redirect_course_management(course.slug)

    context = {
        'chapter': chapter,
        'course': course,
    }
    return render(request, 'lms/admin/confirm_delete_chapter.html', context)


@admin_required
def admin_edit_assignment(request, assignment_id):
    """Edit an existing assignment."""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    course = assignment.course

    if request.method == 'POST':
        rubric_raw = request.POST.get('rubric', '{}')
        try:
            rubric_value = json.loads(rubric_raw)
        except json.JSONDecodeError:
            messages.error(request, "Rubric JSON is invalid.")
            return redirect('lms:admin_edit_assignment', assignment_id=assignment.id)

        assignment.title = request.POST.get('title', '').strip()
        assignment.description = request.POST.get('description', '')
        assignment.assignment_type = request.POST.get('assignment_type', assignment.assignment_type)
        assignment.due_date = request.POST.get('due_date')
        assignment.soft_deadline = request.POST.get('soft_deadline') or None
        assignment.late_penalty_percent = int(request.POST.get('late_penalty_percent', 10))
        assignment.total_points = int(request.POST.get('total_points', 100))
        assignment.rubric = rubric_value
        assignment.enable_plagiarism_check = request.POST.get('enable_plagiarism_check') == 'on'
        assignment.max_file_size_mb = int(request.POST.get('max_file_size_mb', 50))
        assignment.allowed_file_types = request.POST.get('allowed_file_types', '.pdf,.doc,.docx,.txt')
        assignment.save()

        log_action(request.user, 'update', assignment, request)
        messages.success(request, f"Assignment '{assignment.title}' updated successfully.")
        return _redirect_course_management(course.slug)

    context = {
        'assignment': assignment,
        'course': course,
        'assignment_types': Assignment.ASSIGNMENT_TYPE_CHOICES,
    }
    return render(request, 'lms/admin/edit_assignment.html', context)


@admin_required
def admin_delete_assignment(request, assignment_id):
    """Delete an assignment."""
    assignment = get_object_or_404(Assignment, id=assignment_id)
    course = assignment.course
    assignment_title = assignment.title

    if request.method == 'POST':
        log_action(request.user, 'delete', assignment, request)
        assignment.delete()
        messages.success(request, f"Assignment '{assignment_title}' deleted successfully.")
        return _redirect_course_management(course.slug)

    context = {
        'assignment': assignment,
        'course': course,
    }
    return render(request, 'lms/admin/confirm_delete_assignment.html', context)


@admin_required
def admin_edit_quiz(request, quiz_id):
    """Edit an existing quiz."""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    chapter = quiz.chapter
    course = chapter.course

    if request.method == 'POST':
        quiz.title = request.POST.get('title', '').strip() or quiz.title
        quiz.description = request.POST.get('description', '')
        quiz.pass_score = int(request.POST.get('pass_score', 70))
        quiz.time_limit_minutes = int(request.POST.get('time_limit_minutes')) if request.POST.get('time_limit_minutes') else None
        quiz.attempts_allowed = int(request.POST.get('attempts_allowed', 1))
        quiz.shuffle_questions = request.POST.get('shuffle_questions') == 'on'
        quiz.show_answers_after_submit = request.POST.get('show_answers_after_submit') == 'on'
        quiz.save()

        log_action(request.user, 'update', quiz, request)
        messages.success(request, "Quiz updated successfully.")
        return _redirect_course_management(course.slug)

    context = {
        'quiz': quiz,
        'chapter': chapter,
        'course': course,
    }
    return render(request, 'lms/admin/edit_quiz.html', context)


@admin_required
def admin_delete_quiz(request, quiz_id):
    """Delete a quiz."""
    quiz = get_object_or_404(Quiz, id=quiz_id)
    chapter = quiz.chapter
    course = chapter.course

    if request.method == 'POST':
        log_action(request.user, 'delete', quiz, request)
        quiz.delete()
        messages.success(request, "Quiz deleted successfully.")
        return _redirect_course_management(course.slug)

    context = {
        'quiz': quiz,
        'chapter': chapter,
        'course': course,
    }
    return render(request, 'lms/admin/confirm_delete_quiz.html', context)


@admin_required
def admin_edit_question(request, question_id):
    """Edit a quiz question."""
    question = get_object_or_404(Question, id=question_id)
    quiz = question.quiz

    if request.method == 'POST':
        question.text = request.POST.get('question_text', '').strip()
        question.question_type = request.POST.get('question_type', question.question_type)
        question.points = int(request.POST.get('points', 1))

        if question.question_type == 'mcq':
            question.options = [o.strip() for o in request.POST.getlist('options[]') if o.strip()]
            question.correct_answer = request.POST.get('correct_answer', '').strip()
        elif question.question_type == 'tf':
            question.options = ['True', 'False']
            question.correct_answer = request.POST.get('correct_answer', 'True')
        elif question.question_type == 'short':
            question.options = []
            question.correct_answer = request.POST.get('correct_answer', '').strip()

        question.save()
        log_action(request.user, 'update', question, request)
        messages.success(request, "Question updated successfully.")
        return redirect('lms:admin_add_quiz_questions', quiz_id=quiz.id)

    context = {
        'question': question,
        'quiz': quiz,
    }
    return render(request, 'lms/admin/edit_question.html', context)


@admin_required
def admin_delete_question(request, question_id):
    """Delete a quiz question."""
    question = get_object_or_404(Question, id=question_id)
    quiz = question.quiz

    if request.method == 'POST':
        log_action(request.user, 'delete', question, request)
        question.delete()
        messages.success(request, "Question deleted successfully.")
        return redirect('lms:admin_add_quiz_questions', quiz_id=quiz.id)

    context = {
        'question': question,
        'quiz': quiz,
    }
    return render(request, 'lms/admin/confirm_delete_question.html', context)

# -------------------------------------------------------------------
# Admin Views
# -------------------------------------------------------------------

@admin_required
def admin_users(request):
    """User management for admin"""
    users = User.objects.all().order_by('-date_joined')
    
    # Filtering
    role = request.GET.get('role')
    if role:
        users = users.filter(role=role)
    
    status = request.GET.get('status')
    if status == 'approved':
        users = users.filter(is_approved=True)
    elif status == 'pending':
        users = users.filter(is_approved=False)
    
    paginator = Paginator(users, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'role_filter': role,
        'status_filter': status,
        'role_choices': User.ROLE_CHOICES,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    
    # If AJAX request, use minimal template
    if request.GET.get('ajax') == 'true':
        return render(request, 'lms/admin/_user_management_ajax.html', context)
    else:
        return render(request, 'lms/admin/user_management.html', context)

@admin_required
def admin_add_student(request):
    """Add a student manually (direct addition)"""
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        phone_number = request.POST.get('phone_number')
        module_id = request.POST.get('module_id')
        course_ids = request.POST.getlist('course_ids')
        
        # Generate random password
        temp_password = User.objects.make_random_password()
        
        user = User.objects.create_user(
            username=username,
            email=email,
            password=temp_password,
            first_name=first_name,
            last_name=last_name,
            phone_number=phone_number,
            role='learner',
            is_approved=True,
            is_active=True,
            temp_password=temp_password,
            department_id=module_id or None
        )
        
        # Generate approval token for password setup
        token = user.generate_approval_token()
        
        # Enroll in selected courses
        for course_id in course_ids:
            Enrollment.objects.create(
                student=user,
                course_id=course_id,
                status='active',
                enrollment_method='admin',
                enrolled_by=request.user
            )
        
        # Send credentials email
        set_password_url = request.build_absolute_uri(reverse('lms:set_password', args=[token]))
        
        email_html = render_to_string('lms/registration/credentials_email.html', {
            'user': user,
            'temp_password': temp_password,
            'set_password_url': set_password_url
        })
        
        send_email_notification(
            f"Welcome to Synego Institute, {first_name}!",
            f"Your account has been created. Please set your password using: {set_password_url}",
            [email],
            email_html
        )
        
        log_action(request.user, 'create', user, request)
        messages.success(request, f"Student {username} added successfully. Credentials emailed.")
        return redirect('lms:admin_users')
    
    departments = Module.objects.filter(status='active')
    courses = Course.objects.filter(status='published', is_active=True)
    
    context = {
        'modules': departments,
        'courses': courses,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/admin/add_student.html', context)

@admin_required
def admin_enrollment_requests(request):
    """Manage enrollment requests"""
    requests = EnrollmentRequest.objects.select_related('student', 'course').filter(status='pending')
    
    if request.method == 'POST':
        request_id = request.POST.get('request_id')
        action = request.POST.get('action')
        enrollment_request = get_object_or_404(EnrollmentRequest, id=request_id)
        
        if action == 'approve':
            enrollment_request.approve(request.user)
            messages.success(request, f"Enrollment request for {enrollment_request.student.username} approved.")
        elif action == 'reject':
            enrollment_request.reject(request.user, request.POST.get('notes', ''))
            messages.success(request, f"Enrollment request for {enrollment_request.student.username} rejected.")
        
        return redirect('lms:admin_enrollment_requests')
    
    context = {'requests': requests, 'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0}
    return render(request, 'lms/admin/enrollment_requests.html', context)

@admin_required
def admin_reports(request):
    """Generate admin reports"""
    from django.db.models.functions import TruncMonth
    
    report_type = request.GET.get('type', 'enrollment')
    
    if report_type == 'enrollment':
        # Enrollment report
        data = Enrollment.objects.select_related('student', 'course').values(
            'course__title', 'course__department__name'
        ).annotate(
            count=Count('id'),
            completed=Count('id', filter=Q(status='completed'))
        )
    elif report_type == 'completion':
        # Completion report - group by month of issued_at
        data = Certificate.objects.select_related('student', 'course').annotate(
            month=TruncMonth('issued_at')
        ).values('course__title', 'month').annotate(count=Count('id')).order_by('-month')
    else:
        data = []
    
    # Summary statistics
    total_active_students = User.objects.filter(role='learner', is_active=True).count()
    total_enrolled = Enrollment.objects.count()
    total_certificates = Certificate.objects.count()
    total_courses = Course.objects.filter(is_active=True).count()
    
    context = {
        'report_type': report_type,
        'data': data,
        'report_choices': [
            ('enrollment', 'Enrollment Report'),
            ('completion', 'Completion Report'),
            ('quiz', 'Quiz Performance Report'),
            ('assignment', 'Assignment Report'),
        ],
        'total_active_students': total_active_students,
        'total_enrolled': total_enrolled,
        'total_certificates': total_certificates,
        'total_courses': total_courses,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    
    # If AJAX request, use minimal template
    if request.GET.get('ajax') == 'true':
        return render(request, 'lms/admin/_reports_ajax.html', context)
    else:
        return render(request, 'lms/admin/reports.html', context)


@admin_required
def financial_reports(request):
    """Generate financial reports"""
    report_type = request.GET.get('type', 'revenue')
    
    # Generate financial data based on report type
    if report_type == 'revenue':
        # Revenue report data
        data = Enrollment.objects.select_related('course').values(
            'course__title'
        ).annotate(
            total_revenue=Sum('course__price'),
            enrollment_count=Count('id')
        ).order_by('-total_revenue')
    elif report_type == 'expenses':
        # Expenses report (placeholder)
        data = []
    elif report_type == 'profit':
        # Profit & Loss report (placeholder)
        data = []
    elif report_type == 'payments':
        # Student payments report (placeholder)
        data = []
    elif report_type == 'invoices':
        # Invoices report (placeholder)
        data = []
    elif report_type == 'refunds':
        # Refunds report (placeholder)
        data = []
    else:
        data = []
    
    # Financial summary statistics
    total_enrolled = Enrollment.objects.count()
    total_revenue = Enrollment.objects.aggregate(
        total=Sum('course__price')
    )['total'] or 0
    
    context = {
        'report_type': report_type,
        'data': data,
        'total_enrolled': total_enrolled,
        'total_revenue': total_revenue,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    
    return render(request, 'lms/admin/financial_reports.html', context)


@admin_required
def download_report(request, report_type, file_format='csv'):
    """Download report in various formats"""
    
    if file_format == 'csv':
        return download_report_csv(request, report_type)
    elif file_format == 'pdf':
        return download_report_pdf(request, report_type)
    elif file_format == 'excel':
        return download_report_excel(request, report_type)
    elif file_format == 'word':
        return download_report_word(request, report_type)
    else:
        return download_report_csv(request, report_type)

def save_report_to_storage(report_content, report_type, file_format, user):
    """Helper to save generated reports to DocumentStorage"""
    try:
        import io
        from django.core.files.base import ContentFile
        
        filename = f"{report_type}_report_{timezone.now().date()}.{file_format}"
        doc_type = 'report'
        
        # Create a file object
        if isinstance(report_content, str):
            file_obj = ContentFile(report_content.encode('utf-8'), name=filename)
        else:
            file_obj = ContentFile(report_content, name=filename)
        
        # Save to DocumentStorage
        document = DocumentStorage.objects.create(
            name=f"{report_type.replace('_', ' ').title()} Report - {timezone.now().date()}",
            description=f"Auto-generated {report_type} report in {file_format.upper()} format",
            document_type=doc_type,
            file=file_obj,
            uploaded_by=user,
            is_public=False,
            can_download=True,
            can_view_online=True,
        )
        return document
    except Exception as e:
        print(f"Error saving report to storage: {str(e)}")
        return None


def download_report_csv(request, report_type):
    """Download report as CSV - with auto-save to DocumentStorage"""
    import io
    
    # Generate content into StringIO
    csv_buffer = io.StringIO()
    writer = csv.writer(csv_buffer)
    
    if report_type == 'enrollment':
        writer.writerow(['Course', 'Module', 'Total Enrollments', 'Completed'])
        data = Enrollment.objects.select_related('course').values(
            'course__title', 'course__department__name'
        ).annotate(
            total=Count('id'),
            completed=Count('id', filter=Q(status='completed'))
        )
        for row in data:
            writer.writerow([row['course__title'], row['course__department__name'], row['total'], row['completed']])
    
    elif report_type == 'completion':
        writer.writerow(['Course', 'Month', 'Certificates Issued'])
        data = Certificate.objects.select_related('course').values_list('course__title', 'issued_at__month').annotate(count=Count('id'))
        for row in data:
            writer.writerow(row)
    
    elif report_type == 'quiz':
        writer.writerow(['Quiz', 'Average Score', 'Attempts'])
        writer.writerow(['Quiz 1', '78%', '120'])
        writer.writerow(['Quiz 2', '85%', '95'])
    
    elif report_type == 'assignment':
        writer.writerow(['Assignment', 'Course', 'Submissions', 'Graded'])
        writer.writerow(['Assignment 1', 'Python Course', '85', '80'])
        writer.writerow(['Assignment 2', 'Web Dev', '92', '88'])
    
    # Get CSV content
    csv_content = csv_buffer.getvalue()
    csv_buffer.close()
    
    # Auto-save to DocumentStorage
    if request.user.is_authenticated:
        save_report_to_storage(csv_content, report_type, 'csv', request.user)
    
    # Return as HTTP response
    response = HttpResponse(csv_content, content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_report_{timezone.now().date()}.csv"'
    return response


def generate_chart_image(chart_data, chart_type='bar', title='Chart'):
    """Generate a chart image using matplotlib"""
    try:
        import matplotlib
        matplotlib.use('Agg')  # Use non-GUI backend
        import matplotlib.pyplot as plt
        from io import BytesIO
        import base64
        
        fig, ax = plt.subplots(figsize=(8, 5))
        fig.patch.set_facecolor('white')
        
        if chart_type == 'bar':
            labels = chart_data.get('labels', [])
            values = chart_data.get('values', [])
            colors_list = ['#2f80ed', '#50c878', '#f54242', '#ffa500', '#9b59b6']
            ax.bar(labels, values, color=colors_list[:len(labels)], edgecolor='black', linewidth=1.5)
            ax.set_ylabel('Count', fontsize=11, fontweight='bold')
            
        elif chart_type == 'line':
            labels = chart_data.get('labels', [])
            values = chart_data.get('values', [])
            ax.plot(labels, values, marker='o', linewidth=2.5, markersize=8, color='#2f80ed')
            ax.fill_between(range(len(labels)), values, alpha=0.3, color='#2f80ed')
            ax.set_ylabel('Count', fontsize=11, fontweight='bold')
            
        elif chart_type == 'pie':
            labels = chart_data.get('labels', [])
            values = chart_data.get('values', [])
            colors_list = ['#2f80ed', '#50c878', '#f54242', '#ffa500', '#9b59b6', '#1abc9c']
            ax.pie(values, labels=labels, autopct='%1.1f%%', colors=colors_list[:len(labels)],
                   startangle=90, textprops={'fontsize': 9, 'fontweight': 'bold'})
            ax.axis('equal')
        
        ax.set_title(title, fontsize=13, fontweight='bold', color='#1f3a5d', pad=20)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        plt.tight_layout()
        
        # Save to bytes
        img_bytes = BytesIO()
        plt.savefig(img_bytes, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        img_bytes.seek(0)
        plt.close(fig)
        
        return img_bytes
    except Exception as e:
        return None


def download_report_pdf(request, report_type):
    """Download report as PDF with charts and diagrams - with auto-save to DocumentStorage"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import io
        
        # Create BytesIO buffer instead of HttpResponse
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f3a5d'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        title = Paragraph(f'{report_type.title()} Report', title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.3*inch))
        
        # Report Date
        date_para = Paragraph(f'Generated on: {timezone.now().strftime("%B %d, %Y")}', styles['Normal'])
        elements.append(date_para)
        elements.append(Spacer(1, 0.4*inch))
        
        # Get data and create charts
        chart_data = None
        table_data = None
        
        if report_type == 'enrollment':
            enrollments = Enrollment.objects.select_related('course').values(
                'course__title', 'course__department__name'
            ).annotate(
                total=Count('id'),
                completed=Count('id', filter=Q(status='completed'))
            )
            
            table_data = [['Course', 'Module', 'Total Enrollments', 'Completed']]
            courses = []
            totals = []
            for row in enrollments:
                table_data.append([row['course__title'], row['course__department__name'], str(row['total']), str(row['completed'])])
                courses.append(row['course__title'][:15])  # Truncate for chart
                totals.append(row['total'])
            
            if courses:
                chart_data = {'labels': courses, 'values': totals}
                # Add bar chart
                chart_img = generate_chart_image(chart_data, 'bar', 'Total Enrollments by Course')
                if chart_img:
                    elements.append(RLImage(chart_img, width=6*inch, height=3.5*inch))
                    elements.append(Spacer(1, 0.3*inch))
        
        elif report_type == 'completion':
            certificates = Certificate.objects.select_related('course').values_list(
                'course__title', 'issued_at__month'
            ).annotate(count=Count('id')).order_by('issued_at__month')
            
            table_data = [['Course', 'Month', 'Certificates Issued']]
            months = []
            counts = []
            for row in certificates:
                table_data.append(list(row))
                month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
                if row[1]:
                    months.append(month_names[row[1]-1])
                    counts.append(row[2])
            
            if months:
                chart_data = {'labels': months, 'values': counts}
                chart_img = generate_chart_image(chart_data, 'line', 'Certificates Issued Over Time')
                if chart_img:
                    elements.append(RLImage(chart_img, width=6*inch, height=3.5*inch))
                    elements.append(Spacer(1, 0.3*inch))
        
        elif report_type == 'quiz':
            # Sample quiz data with chart
            quiz_labels = ['Quiz 1', 'Quiz 2', 'Quiz 3', 'Quiz 4']
            quiz_scores = [78, 85, 82, 90]
            chart_data = {'labels': quiz_labels, 'values': quiz_scores}
            chart_img = generate_chart_image(chart_data, 'bar', 'Average Quiz Scores')
            if chart_img:
                elements.append(RLImage(chart_img, width=6*inch, height=3.5*inch))
                elements.append(Spacer(1, 0.3*inch))
            
            table_data = [['Quiz', 'Average Score', 'Attempts']]
            table_data.append(['Quiz 1', '78%', '120'])
            table_data.append(['Quiz 2', '85%', '95'])
            table_data.append(['Quiz 3', '82%', '110'])
            table_data.append(['Quiz 4', '90%', '85'])
        
        elif report_type == 'assignment':
            # Sample assignment data with chart
            assign_labels = ['Assignment 1', 'Assignment 2', 'Assignment 3']
            assign_values = [85, 92, 78]
            chart_data = {'labels': assign_labels, 'values': assign_values}
            chart_img = generate_chart_image(chart_data, 'bar', 'Assignment Submission Rates')
            if chart_img:
                elements.append(RLImage(chart_img, width=6*inch, height=3.5*inch))
                elements.append(Spacer(1, 0.3*inch))
            
            table_data = [['Assignment', 'Course', 'Submissions', 'Graded']]
            table_data.append(['Assignment 1', 'Python Course', '85', '80'])
            table_data.append(['Assignment 2', 'Web Dev', '92', '88'])
            table_data.append(['Assignment 3', 'Data Science', '78', '75'])
        
        else:
            table_data = [['Item', 'Value'], ['Sample Data', f'{report_type} information']]
        
        # Add section heading for table
        table_heading = Paragraph('<b>Detailed Data Table</b>', styles['Heading2'])
        elements.append(table_heading)
        elements.append(Spacer(1, 0.2*inch))
        
        # Create and style table
        if table_data:
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2f80ed')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f5ff')])
            ]))
            elements.append(table)
        
        # Build PDF to buffer
        doc.build(elements)
        pdf_buffer.seek(0)
        pdf_content = pdf_buffer.getvalue()
        
        # Auto-save to DocumentStorage
        if request.user.is_authenticated:
            save_report_to_storage(pdf_content, report_type, 'pdf', request.user)
        
        # Return as HTTP response
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{report_type}_report_{timezone.now().date()}.pdf"'
        return response
    except ImportError as e:
        return HttpResponse(f'Required library missing: {str(e)}. Please install with: pip install matplotlib reportlab', status=500)


def download_report_excel(request, report_type):
    """Download report as Excel with charts"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.drawing.image import Image as XLImage
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{report_type}_report_{timezone.now().date()}.xlsx"'
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = report_type.title()
        
        # Title
        ws['A1'] = f'{report_type.title()} Report'
        ws['A1'].font = Font(size=16, bold=True, color='1F3A5D')
        ws['A1'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A1:D1')
        
        # Date
        ws['A2'] = f'Generated: {timezone.now().strftime("%B %d, %Y")}'
        ws['A2'].font = Font(size=10, italic=True)
        ws.merge_cells('A2:D2')
        
        # Headers and data
        current_row = 4
        headers = []
        rows_data = []
        chart_img = None
        chart_type = 'bar'
        
        if report_type == 'enrollment':
            headers = ['Course', 'Module', 'Total Enrollments', 'Completed']
            enrollments = Enrollment.objects.select_related('course').values(
                'course__title', 'course__department__name'
            ).annotate(
                total=Count('id'),
                completed=Count('id', filter=Q(status='completed'))
            )
            for row in enrollments:
                rows_data.append([row['course__title'], row['course__department__name'], row['total'], row['completed']])
            
            # Generate chart data
            courses = [r[0][:15] for r in rows_data]
            totals = [r[2] for r in rows_data]
            if courses:
                chart_data = {'labels': courses, 'values': totals}
                chart_img = generate_chart_image(chart_data, 'bar', 'Total Enrollments by Course')
        
        elif report_type == 'completion':
            headers = ['Course', 'Month', 'Certificates Issued']
            certificates = Certificate.objects.select_related('course').values_list(
                'course__title', 'issued_at__month'
            ).annotate(count=Count('id')).order_by('issued_at__month')
            for row in certificates:
                rows_data.append(list(row))
            
            # Generate chart data
            months = []
            counts = []
            month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
            for row in rows_data:
                if row[1]:
                    months.append(month_names[row[1]-1])
                    counts.append(row[2])
            if months:
                chart_data = {'labels': months, 'values': counts}
                chart_img = generate_chart_image(chart_data, 'line', 'Certificates Issued Over Time')
                chart_type = 'line'
        
        elif report_type == 'quiz':
            headers = ['Quiz', 'Average Score', 'Attempts']
            rows_data = [
                ['Quiz 1', '78%', '120'],
                ['Quiz 2', '85%', '95'],
                ['Quiz 3', '82%', '110'],
                ['Quiz 4', '90%', '85']
            ]
            chart_data = {'labels': [r[0] for r in rows_data], 'values': [int(r[1].rstrip('%')) for r in rows_data]}
            chart_img = generate_chart_image(chart_data, 'bar', 'Average Quiz Scores')
        
        elif report_type == 'assignment':
            headers = ['Assignment', 'Course', 'Submissions', 'Graded']
            rows_data = [
                ['Assignment 1', 'Python Course', '85', '80'],
                ['Assignment 2', 'Web Dev', '92', '88'],
                ['Assignment 3', 'Data Science', '78', '75']
            ]
            chart_data = {'labels': [r[0] for r in rows_data], 'values': [int(r[2]) for r in rows_data]}
            chart_img = generate_chart_image(chart_data, 'bar', 'Assignment Submission Rates')
        
        else:
            headers = ['Item', 'Value']
            rows_data = [['Report Type', report_type], ['Generated Date', timezone.now().strftime('%Y-%m-%d')]]
        
        # Add chart if available
        if chart_img and headers:
            ws.add_row([''])  # Blank row
            try:
                xl_image = XLImage(chart_img)
                xl_image.width = 600
                xl_image.height = 350
                ws.add_image(xl_image, 'A5')
                current_row = 22  # Move down to make room for chart
            except:
                pass  # If chart embedding fails, continue with just data
        
        # Add headers row
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=current_row, column=col, value=header)
            cell.fill = PatternFill(start_color='2F80ED', end_color='2F80ED', fill_type='solid')
            cell.font = Font(bold=True, color='FFFFFF')
            cell.alignment = Alignment(horizontal='center')
        
        # Add data rows
        for row_idx, row_data in enumerate(rows_data, start=current_row+1):
            for col_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.alignment = Alignment(horizontal='center')
                if row_idx % 2 == 0:
                    cell.fill = PatternFill(start_color='F0F5FF', end_color='F0F5FF', fill_type='solid')
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 20
        
        # Save workbook to buffer
        import io
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)
        excel_content = excel_buffer.getvalue()
        
        # Auto-save to DocumentStorage
        if request.user.is_authenticated:
            save_report_to_storage(excel_content, report_type, 'xlsx', request.user)
        
        # Return as HTTP response
        response = HttpResponse(excel_content, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{report_type}_report_{timezone.now().date()}.xlsx"'
        return response
    except ImportError as e:
        return HttpResponse(f'openpyxl is not installed. Please install it with: pip install openpyxl', status=500)


def download_report_word(request, report_type):
    """Download report as Word document with charts"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{report_type}_report_{timezone.now().date()}.docx"'
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{report_type.title()} Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(31, 58, 93)
        
        # Date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Generated on: {timezone.now().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(10)
        date_run.italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Blank line
        
        # Add chart if available
        chart_img = None
        chart_title = ''
        enrollments_data = []
        certificates_data = []
        
        if report_type == 'enrollment':
            enrollments_data = list(Enrollment.objects.select_related('course').values(
                'course__title', 'course__department__name'
            ).annotate(
                total=Count('id'),
                completed=Count('id', filter=Q(status='completed'))
            ))
            
            courses = [r['course__title'][:15] for r in enrollments_data]
            totals = [r['total'] for r in enrollments_data]
            if courses:
                chart_data = {'labels': courses, 'values': totals}
                chart_img = generate_chart_image(chart_data, 'bar', 'Total Enrollments by Course')
                chart_title = 'Enrollment Chart'
        
        elif report_type == 'completion':
            certificates_data = list(Certificate.objects.select_related('course').values_list(
                'course__title', 'issued_at__month'
            ).annotate(count=Count('id')).order_by('issued_at__month'))
            
            months = []
            counts = []
            month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
            for row in certificates_data:
                if row[1]:
                    months.append(month_names[row[1]-1])
                    counts.append(row[2])
            if months:
                chart_data = {'labels': months, 'values': counts}
                chart_img = generate_chart_image(chart_data, 'line', 'Certificates Issued Over Time')
                chart_title = 'Certificate Issuance Chart'
        
        elif report_type == 'quiz':
            quiz_data = [
                {'title': 'Quiz 1', 'score': 78, 'attempts': 120},
                {'title': 'Quiz 2', 'score': 85, 'attempts': 95},
                {'title': 'Quiz 3', 'score': 82, 'attempts': 110},
                {'title': 'Quiz 4', 'score': 90, 'attempts': 85}
            ]
            chart_data = {'labels': [q['title'] for q in quiz_data], 'values': [q['score'] for q in quiz_data]}
            chart_img = generate_chart_image(chart_data, 'bar', 'Average Quiz Scores')
            chart_title = 'Quiz Performance Chart'
        
        elif report_type == 'assignment':
            assignment_data = [
                {'name': 'Assignment 1', 'course': 'Python Course', 'submissions': 85, 'graded': 80},
                {'name': 'Assignment 2', 'course': 'Web Dev', 'submissions': 92, 'graded': 88},
                {'name': 'Assignment 3', 'course': 'Data Science', 'submissions': 78, 'graded': 75}
            ]
            chart_data = {'labels': [a['name'] for a in assignment_data], 'values': [a['submissions'] for a in assignment_data]}
            chart_img = generate_chart_image(chart_data, 'bar', 'Assignment Submission Rates')
            chart_title = 'Assignment Submission Chart'
        
        # Add chart heading and image if available
        if chart_img:
            doc.add_heading(chart_title, level=2)
            try:
                doc.add_picture(chart_img, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Blank line
            except:
                pass  # If chart fails, continue with table
        
        # Summary section
        doc.add_heading('Report Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f'Report Type: ')
        summary_run.bold = True
        summary_para.add_run(report_type.title())
        
        # Data table
        doc.add_heading('Detailed Data Table', level=1)
        
        if report_type == 'enrollment' and enrollments_data:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Course', 'Module', 'Total Enrollments', 'Completed']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.style = 'Heading Row 1 Char'
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '2F80ED')
                tcPr.append(tcVAlign)
            
            # Data rows
            for row in enrollments_data:
                row_cells = table.add_row().cells
                row_cells[0].text = row['course__title']
                row_cells[1].text = row['course__department__name']
                row_cells[2].text = str(row['total'])
                row_cells[3].text = str(row['completed'])
        
        elif report_type == 'completion' and certificates_data:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Course', 'Month', 'Certificates Issued']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '2F80ED')
                tcPr.append(tcVAlign)
            
            # Data rows
            month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
            for row in certificates_data:
                row_cells = table.add_row().cells
                row_cells[0].text = row[0]
                row_cells[1].text = month_names[row[1]-1] if row[1] else 'N/A'
                row_cells[2].text = str(row[2])
        
        elif report_type == 'quiz':
            quiz_data = [
                ['Quiz 1', '78%', '120'],
                ['Quiz 2', '85%', '95'],
                ['Quiz 3', '82%', '110'],
                ['Quiz 4', '90%', '85']
            ]
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Quiz', 'Average Score', 'Attempts']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '2F80ED')
                tcPr.append(tcVAlign)
            
            # Data rows
            for row_data in quiz_data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data):
                    row_cells[i].text = value
        
        elif report_type == 'assignment':
            assignment_data = [
                ['Assignment 1', 'Python Course', '85', '80'],
                ['Assignment 2', 'Web Dev', '92', '88'],
                ['Assignment 3', 'Data Science', '78', '75']
            ]
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Assignment', 'Course', 'Submissions', 'Graded']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                tcPr = header_cells[i]._element.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), '2F80ED')
                tcPr.append(tcVAlign)
            
            # Data rows
            for row_data in assignment_data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data):
                    row_cells[i].text = value
        
        else:
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Light Grid Accent 1'
            table.rows[0].cells[0].text = 'Item'
            table.rows[0].cells[1].text = 'Value'
            table.rows[1].cells[0].text = 'Report Type'
            table.rows[1].cells[1].text = report_type.title()
            table.rows[2].cells[0].text = 'Generated Date'
            table.rows[2].cells[1].text = timezone.now().strftime('%Y-%m-%d')
        
        # Save document to buffer
        import io
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        word_content = word_buffer.getvalue()
        
        # Auto-save to DocumentStorage
        if request.user.is_authenticated:
            save_report_to_storage(word_content, report_type, 'docx', request.user)
        
        # Return as HTTP response
        response = HttpResponse(word_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{report_type}_report_{timezone.now().date()}.docx"'
        return response
    except ImportError as e:
        return HttpResponse(f'python-docx is not installed. Please install it with: pip install python-docx', status=500)


@admin_required
def admin_system_logs(request):
    """View system logs for admin"""
    logs = SystemLog.objects.select_related('user').all().order_by('-created_at')
    
    # Filtering
    action = request.GET.get('action')
    if action:
        logs = logs.filter(action=action)
    
    object_type = request.GET.get('object_type')
    if object_type:
        logs = logs.filter(object_type=object_type)
    
    user_id = request.GET.get('user')
    if user_id:
        logs = logs.filter(user_id=user_id)
    
    # Pagination
    paginator = Paginator(logs, 50)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Get filter options
    actions = SystemLog.objects.values_list('action', flat=True).distinct()
    object_types = SystemLog.objects.values_list('object_type', flat=True).distinct()
    users = User.objects.filter(id__in=SystemLog.objects.values_list('user_id', flat=True).distinct())
    
    context = {
        'page_obj': page_obj,
        'actions': actions,
        'object_types': object_types,
        'users': users,
        'selected_action': action,
        'selected_object_type': object_type,
        'selected_user': user_id,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/admin/system_logs.html', context)


@admin_required
def admin_edit_user(request, user_id):
    """Edit user from admin panel"""
    edit_user = get_object_or_404(User, id=user_id)
    
    if request.method == 'POST':
        edit_user.username = request.POST.get('username')
        edit_user.email = request.POST.get('email')
        edit_user.first_name = request.POST.get('first_name')
        edit_user.last_name = request.POST.get('last_name')
        edit_user.phone_number = request.POST.get('phone_number')
        edit_user.role = request.POST.get('role')
        edit_user.module_id = request.POST.get('module_id') or None
        edit_user.is_approved = request.POST.get('is_approved') == 'on'
        edit_user.is_active = request.POST.get('is_active') == 'on'
        
        edit_user.save()
        
        log_action(request.user, 'update', edit_user, request)
        messages.success(request, f"User {edit_user.username} updated successfully.")
        return redirect('lms:admin_users')
    
    departments = Module.objects.filter(status='active')
    context = {
        'edit_user': edit_user,
        'modules': departments,
        'role_choices': User.ROLE_CHOICES,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/admin/edit_user.html', context)


@admin_required
def admin_delete_user(request, user_id):
    """Delete user from admin panel"""
    delete_user = get_object_or_404(User, id=user_id)
    
    if request.method == 'POST':
        username = delete_user.username
        delete_user.delete()
        log_action(request.user, 'delete', delete_user, request)
        messages.success(request, f"User {username} deleted successfully.")
        return redirect('lms:admin_users')
    
    return render(request, 'lms/admin/confirm_delete_user.html', {'user': delete_user, 'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0})


@admin_required
def admin_modules(request):
    """Manage modules"""
    modules = Module.objects.all().order_by('name')
    
    context = {
        'modules': modules,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    
    # If AJAX request, use minimal template
    if request.GET.get('ajax') == 'true':
        return render(request, 'lms/admin/_modules_ajax.html', context)
    else:
        return render(request, 'lms/admin/modules.html', context)


@admin_required
def admin_add_module(request):
    """Add a new module"""
    if request.method == 'POST':
        name = request.POST.get('name')
        code = request.POST.get('code')
        description = request.POST.get('description')
        mission = request.POST.get('mission')
        vision = request.POST.get('vision')
        infrastructure = request.POST.get('infrastructure')
        resources = request.POST.get('resources')
        contact_email = request.POST.get('contact_email')
        contact_phone = request.POST.get('contact_phone')
        office_location = request.POST.get('office_location')
        min_instructors = request.POST.get('min_instructors', 1)
        max_capacity = request.POST.get('max_capacity', 100)
        
        module = Module.objects.create(
            name=name,
            code=code,
            description=description,
            mission=mission,
            vision=vision,
            infrastructure=infrastructure,
            resources=resources,
            contact_email=contact_email,
            contact_phone=contact_phone,
            office_location=office_location,
            min_instructors=min_instructors,
            max_capacity=max_capacity
        )
        
        log_action(request.user, 'create', module, request)
        messages.success(request, f"Module {name} created successfully.")
        return redirect('lms:admin_modules')
    
    return render(request, 'lms/admin/add_module.html', {'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0})


@admin_required
def admin_edit_module(request, module_id):
    """Edit a module"""
    module = get_object_or_404(Module, id=module_id)
    
    if request.method == 'POST':
        module.name = request.POST.get('name')
        module.code = request.POST.get('code')
        module.description = request.POST.get('description')
        module.mission = request.POST.get('mission')
        module.vision = request.POST.get('vision')
        module.infrastructure = request.POST.get('infrastructure')
        module.resources = request.POST.get('resources')
        module.contact_email = request.POST.get('contact_email')
        module.contact_phone = request.POST.get('contact_phone')
        module.office_location = request.POST.get('office_location')
        module.min_instructors = request.POST.get('min_instructors', 1)
        module.max_capacity = request.POST.get('max_capacity', 100)
        module.status = request.POST.get('status')
        
        module.save()
        
        log_action(request.user, 'update', module, request)
        messages.success(request, f"Module {module.name} updated successfully.")
        return redirect('lms:admin_modules')
    
    context = {'module': module, 'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0}
    return render(request, 'lms/admin/edit_module.html', context)


@admin_required
def admin_process_request(request, request_id):
    """Process an enrollment request via AJAX or direct POST"""
    enrollment_request = get_object_or_404(EnrollmentRequest, id=request_id)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        notes = request.POST.get('notes', '')
        
        if action == 'approve':
            enrollment_request.approve(request.user)
            messages.success(request, f"Enrollment request for {enrollment_request.student.username} approved.")
        elif action == 'reject':
            enrollment_request.reject(request.user, notes)
            messages.success(request, f"Enrollment request for {enrollment_request.student.username} rejected.")
        
        return redirect('lms:admin_enrollment_requests')
    
    context = {'request_obj': enrollment_request, 'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0}
    return render(request, 'lms/admin/process_request.html', context)


@admin_required
def download_bulk_enroll_template(request):
    """Download CSV template for bulk enrollment"""
    import csv
    from django.http import HttpResponse
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="bulk_enrollment_template.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['email', 'username', 'first_name', 'last_name'])
    writer.writerow(['student1@example.com', 'student1', 'John', 'Doe'])
    writer.writerow(['student2@example.com', 'student2', 'Jane', 'Smith'])
    writer.writerow(['student3@example.com', 'student3', 'Bob', 'Johnson'])
    
    return response


def download_bulk_courses_template(request):
    """Download CSV template for bulk course upload"""
    import csv
    from django.http import HttpResponse
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="bulk_courses_template.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['title', 'code', 'description', 'department_id', 'instructor_id', 'capacity'])
    writer.writerow(['Introduction to Python', 'COMP101', 'Learn Python basics', '1', '5', '40'])
    writer.writerow(['Advanced JavaScript', 'WEB201', 'Master modern JavaScript', '2', '6', '35'])
    writer.writerow(['Data Science 101', 'DATA301', 'Data analysis and visualization', '3', '7', '30'])
    
    return response


@admin_required
def bulk_enroll_students(request):
    """Bulk enroll students via CSV upload"""
    if request.method == 'POST':
        csv_file = request.FILES.get('csv_file')
        course_id = request.POST.get('course_id')
        
        if not csv_file or not course_id:
            messages.error(request, "Please provide both CSV file and course.")
            return redirect('lms:bulk_enroll_students')
        
        course = get_object_or_404(Course, id=course_id)
        
        # Read CSV
        decoded_file = csv_file.read().decode('utf-8')
        io_string = io.StringIO(decoded_file)
        reader = csv.DictReader(io_string)
        
        created_count = 0
        enrolled_count = 0
        errors = []
        
        for row in reader:
            email = row.get('email')
            username = row.get('username', email.split('@')[0])
            first_name = row.get('first_name', '')
            last_name = row.get('last_name', '')
            
            # Get or create user
            user, created = User.objects.get_or_create(
                email=email,
                defaults={
                    'username': username,
                    'first_name': first_name,
                    'last_name': last_name,
                    'role': 'learner',
                    'is_approved': True,
                    'is_active': True
                }
            )
            
            if created:
                # Generate temp password
                temp_password = User.objects.make_random_password()
                user.set_password(temp_password)
                user.temp_password = temp_password
                user.save()
                created_count += 1
            
            # Enroll in course
            enrollment, enrolled = Enrollment.objects.get_or_create(
                student=user,
                course=course,
                defaults={'status': 'active', 'enrollment_method': 'bulk'}
            )
            if enrolled:
                enrolled_count += 1
        
        messages.success(request, f"Bulk enrollment complete! {created_count} users created, {enrolled_count} enrollments added.")
        if errors:
            messages.warning(request, f"Errors: {', '.join(errors[:5])}")
        
        return redirect('lms:admin_users')
    
    courses = Course.objects.filter(status='published', is_active=True)
    context = {'courses': courses, 'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0}
    return render(request, 'lms/admin/bulk_enroll.html', context)


@admin_required
def bulk_upload_courses(request):
    """Bulk upload courses via CSV"""
    if request.method == 'POST':
        csv_file = request.FILES.get('csv_file')
        
        if not csv_file:
            messages.error(request, "Please provide a CSV file.")
            return redirect('lms:bulk_upload_courses')
        
        decoded_file = csv_file.read().decode('utf-8')
        io_string = io.StringIO(decoded_file)
        reader = csv.DictReader(io_string)
        
        created_count = 0
        errors = []
        
        for row in reader:
            try:
                department = Module.objects.get(code=row.get('department_code'))
                
                course = Course.objects.create(
                    department=department,
                    title=row.get('title'),
                    code=row.get('code'),
                    description=row.get('description', ''),
                    duration=row.get('duration'),
                    difficulty=row.get('difficulty', 'beginner'),
                    price=float(row.get('price', 0)),
                    status='draft'
                )
                created_count += 1
            except Exception as e:
                errors.append(f"{row.get('code', 'Unknown')}: {str(e)}")
        
        messages.success(request, f"Bulk upload complete! {created_count} courses created.")
        if errors:
            messages.warning(request, f"Errors: {', '.join(errors[:5])}")
        
        return redirect('lms:course_list')
    
    return render(request, 'lms/admin/bulk_upload_courses.html', {'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0})


# -------------------------------------------------------------------
# Course Management Views (Admin)
# -------------------------------------------------------------------

@admin_required
def admin_create_course(request):
    """Create a new course"""
    if request.method == 'POST':
        title = request.POST.get('title')
        code = request.POST.get('code')
        description = request.POST.get('description')
        department_id = request.POST.get('department_id')
        
        try:
            department = Module.objects.get(id=department_id)
            course = Course.objects.create(
                title=title,
                code=code,
                description=description,
                department=department,
                status='draft'
            )
            messages.success(request, f"Course '{title}' created successfully.")
            return redirect('lms:course_list')
        except Module.DoesNotExist:
            messages.error(request, "Selected department does not exist.")
        except Exception as e:
            messages.error(request, f"Error creating course: {str(e)}")
    
    departments = Module.objects.all()
    context = {
        'departments': departments,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/admin/create_course.html', context)


@admin_required
def admin_edit_course(request, course_slug):
    """Edit an existing course"""
    course = get_object_or_404(Course, slug=course_slug)
    
    if request.method == 'POST':
        course.title = request.POST.get('title', course.title)
        course.code = request.POST.get('code', course.code)
        course.description = request.POST.get('description', course.description)
        department_id = request.POST.get('department_id')
        
        if department_id:
            try:
                course.department = Module.objects.get(id=department_id)
            except Module.DoesNotExist:
                messages.error(request, "Selected department does not exist.")
                return redirect('lms:admin_edit_course', course_slug=course_slug)
        
        try:
            course.save()
            messages.success(request, f"Course '{course.title}' updated successfully.")
            return redirect('lms:course_list')
        except Exception as e:
            messages.error(request, f"Error updating course: {str(e)}")
    
    departments = Module.objects.all()
    context = {
        'course': course,
        'departments': departments,
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
    }
    return render(request, 'lms/admin/edit_course.html', context)

# -------------------------------------------------------------------
# Certificate Views
# -------------------------------------------------------------------

def view_certificate(request, certificate_id):
    """View a certificate"""
    certificate = get_object_or_404(Certificate, certificate_id=certificate_id)
    return render(request, 'lms/certificates/certificate_view.html', {'certificate': certificate})

def download_certificate(request, certificate_id):
    """Download certificate PDF"""
    certificate = get_object_or_404(Certificate, certificate_id=certificate_id)
    response = HttpResponse(certificate.pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="certificate_{certificate_id}.pdf"'
    return response

def verify_certificate(request, certificate_id):
    """Verify certificate authenticity"""
    try:
        certificate = Certificate.objects.get(certificate_id=certificate_id)
        is_valid = certificate.verify()
    except Certificate.DoesNotExist:
        certificate = None
        is_valid = False
    
    context = {
        'certificate': certificate,
        'is_valid': is_valid,
    }
    return render(request, 'lms/certificates/verify_certificate.html', context)

# -------------------------------------------------------------------
# Google Integration Views (Placeholders)
# -------------------------------------------------------------------

@login_required
def google_login(request):
    """Initiate Google OAuth login"""
    # This would integrate with django-allauth
    return redirect('/accounts/google/login/')

@login_required
def google_callback(request):
    """Handle Google OAuth callback"""
    # Handled by allauth
    return redirect('lms:unified_dashboard')

@login_required
def sync_google_classroom(request):
    """Sync courses from Google Classroom"""
    messages.success(request, "Google Classroom sync initiated.")
    return redirect('lms:unified_dashboard')

@login_required
def create_google_doc(request, assignment_id):
    """Create a Google Doc from template"""
    messages.success(request, "Google Doc created.")
    return redirect('lms:assignment_detail', assignment_id=assignment_id)

# -------------------------------------------------------------------
# Turnitin Integration Views (Placeholders)
# -------------------------------------------------------------------

@login_required
def turnitin_launch(request, submission_id):
    """Launch Turnitin LTI"""
    submission = get_object_or_404(Submission, id=submission_id)
    # This would implement LTI 1.3 launch
    return render(request, 'lms/turnitin/turnitin_launch.html', {'submission': submission})

@csrf_exempt
def turnitin_callback(request):
    """Handle Turnitin callback"""
    return JsonResponse({'status': 'ok'})

@login_required
def turnitin_report(request, submission_id):
    """View Turnitin similarity report"""
    submission = get_object_or_404(Submission, id=submission_id)
    return render(request, 'lms/turnitin/turnitin_report.html', {'submission': submission})

# -------------------------------------------------------------------
# Progress & Grades Views
# -------------------------------------------------------------------

@learner_required
def my_progress(request):
    """View overall progress"""
    enrollments = request.user.enrollments.filter(status='active').select_related('course')
    
    for enrollment in enrollments:
        enrollment.progress_percent = enrollment.get_progress_percent()
    
    context = {'enrollments': enrollments}
    return render(request, 'lms/student/my_progress.html', context)

@learner_required
def course_progress(request, course_slug):
    """View progress for a specific course"""
    course = get_object_or_404(Course, slug=course_slug)
    enrollment = get_object_or_404(Enrollment, student=request.user, course=course)
    
    chapters = course.chapters.all().order_by('order')
    completed_chapters = enrollment.progress.get('completed_chapters', [])
    
    for chapter in chapters:
        chapter.is_completed = str(chapter.id) in completed_chapters
    
    context = {
        'course': course,
        'enrollment': enrollment,
        'chapters': chapters,
    }
    return render(request, 'lms/student/course_progress.html', context)

@learner_required
def my_grades(request):
    """View all grades"""
    submissions = Submission.objects.filter(
        student=request.user,
        status='graded',
        grade__isnull=False
    ).select_related('assignment', 'assignment__course')
    
    context = {'submissions': submissions}
    return render(request, 'lms/student/my_grades.html', context)

@learner_required
def course_grades(request, course_slug):
    """View grades for a specific course"""
    course = get_object_or_404(Course, slug=course_slug)
    submissions = Submission.objects.filter(
        student=request.user,
        assignment__course=course,
        status='graded',
        grade__isnull=False
    ).select_related('assignment')
    
    context = {
        'course': course,
        'submissions': submissions,
    }
    return render(request, 'lms/student/course_grades.html', context)

# -------------------------------------------------------------------
# Notification Views
# -------------------------------------------------------------------

@login_required
def notification_list(request):
    """List all notifications"""
    notifications = request.user.notifications.all()
    
    # Mark as read if requested
    if request.GET.get('mark_read'):
        notification_id = request.GET.get('notification_id')
        if notification_id:
            notification = get_object_or_404(Notification, id=notification_id, recipient=request.user)
            notification.mark_as_read()
        else:
            request.user.notifications.filter(is_read=False).update(is_read=True)
    
    paginator = Paginator(notifications, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {'page_obj': page_obj}
    return render(request, 'lms/notifications/notification_list.html', context)

@login_required
def mark_all_notifications_read(request):
    """Mark all notifications as read"""
    request.user.notifications.filter(is_read=False).update(is_read=True, read_at=timezone.now())
    messages.success(request, "All notifications marked as read.")
    return redirect('lms:notification_list')

# -------------------------------------------------------------------
# Search View
# -------------------------------------------------------------------

def search(request):
    """Global search"""
    query = request.GET.get('q', '')
    
    if not query:
        return render(request, 'lms/search.html', {'query': query, 'results': []})
    
    courses = Course.objects.filter(
        Q(title__icontains=query) | Q(description__icontains=query),
        status='published', is_active=True
    )
    
    departments = Module.objects.filter(
        Q(name__icontains=query) | Q(description__icontains=query),
        status='active'
    )
    
    users = User.objects.filter(
        Q(first_name__icontains=query) | Q(last_name__icontains=query) | Q(email__icontains=query),
        is_approved=True
    )[:10] if request.user.is_authenticated and request.user.role in ['admin', 'instructor'] else []
    
    context = {
        'query': query,
        'courses': courses,
        'modules': departments,
        'users': users,
    }
    return render(request, 'lms/search.html', context)

# -------------------------------------------------------------------
# API Endpoints (AJAX)
# -------------------------------------------------------------------

@login_required
def api_chapter_progress(request):
    """Get chapter progress for a course"""
    course_slug = request.GET.get('course_slug')
    course = get_object_or_404(Course, slug=course_slug)
    enrollment = Enrollment.objects.filter(student=request.user, course=course).first()
    
    if not enrollment:
        return JsonResponse({'error': 'Not enrolled'}, status=403)
    
    return JsonResponse({'progress': enrollment.progress})

@login_required
def api_update_progress(request):
    """Update chapter progress via AJAX"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    
    data = json.loads(request.body)
    chapter_id = data.get('chapter_id')
    completed = data.get('completed', False)
    
    chapter = get_object_or_404(Chapter, id=chapter_id)
    enrollment = Enrollment.objects.filter(student=request.user, course=chapter.course).first()
    
    if not enrollment:
        return JsonResponse({'error': 'Not enrolled'}, status=403)
    
    progress = enrollment.progress or {}
    completed_chapters = progress.get('completed_chapters', [])
    
    if completed and str(chapter_id) not in completed_chapters:
        completed_chapters.append(str(chapter_id))
    elif not completed and str(chapter_id) in completed_chapters:
        completed_chapters.remove(str(chapter_id))
    
    progress['completed_chapters'] = completed_chapters
    enrollment.progress = progress
    enrollment.save()
    
    return JsonResponse({'status': 'ok', 'progress_percent': enrollment.get_progress_percent()})

@login_required
def api_mark_notification_read(request):
    """Mark a notification as read via AJAX"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=400)
    
    data = json.loads(request.body)
    notification_id = data.get('notification_id')
    
    notification = get_object_or_404(Notification, id=notification_id, recipient=request.user)
    notification.mark_as_read()
    
    return JsonResponse({'status': 'ok'})

@login_required
def api_get_notifications(request):
    """Get unread notifications via AJAX"""
    notifications = request.user.notifications.filter(is_read=False).values('id', 'title', 'message', 'created_at')
    return JsonResponse({'notifications': list(notifications)})

# -------------------------------------------------------------------
# Help & Support Views
# -------------------------------------------------------------------

def help_center(request):
    """Help center landing page"""
    topics = [
        {'slug': 'getting-started', 'title': 'Getting Started', 'icon': '🚀'},
        {'slug': 'courses', 'title': 'Courses & Enrollment', 'icon': '📚'},
        {'slug': 'assignments', 'title': 'Assignments', 'icon': '📝'},
        {'slug': 'certificates', 'title': 'Certificates', 'icon': '🎓'},
        {'slug': 'technical', 'title': 'Technical Support', 'icon': '💻'},
    ]
    return render(request, 'lms/help/help_center.html', {'topics': topics})

def help_topic(request, topic):
    """Help topic detail"""
    return render(request, f'lms/help/{topic}.html', {'topic': topic})

@login_required
def create_support_ticket(request):
    """Create a support ticket"""
    if request.method == 'POST':
        subject = request.POST.get('subject')
        message = request.POST.get('message')
        
        # This would integrate with a ticketing system
        send_email_notification(
            f"Support Ticket: {subject}",
            f"From: {request.user.email}\n\n{message}",
            [settings.DEFAULT_FROM_EMAIL]
        )
        
        messages.success(request, "Support ticket created. We'll respond within 24 hours.")
        return redirect('lms:help_center')
    
    return render(request, 'lms/help/create_ticket.html')

# -------------------------------------------------------------------
# System Settings & Enterprise Features Dashboard
# -------------------------------------------------------------------

@login_required
@admin_required
def admin_system_settings(request):
    """System settings and enterprise features dashboard - All advanced features have been removed"""
    context = {
        'features': {},
        'unread_notifications': request.user.notifications.filter(is_read=False).count() if request.user.is_authenticated else 0,
        'message': 'The advanced features (AI Administration, Blockchain, Integrations, Security, Configuration, and Analytics) have been removed from this system.'
    }
    
    return render(request, 'lms/admin/system_settings.html', context)

# -------------------------------------------------------------------
# Error Handlers
# -------------------------------------------------------------------

def bad_request(request, exception=None):
    return render(request, 'lms/400.html', status=400)

def permission_denied(request, exception=None):
    return render(request, 'lms/403.html', status=403)

def page_not_found(request, exception=None):
    return render(request, 'lms/404.html', status=404)

def server_error(request):
    return render(request, 'lms/500.html', status=500)